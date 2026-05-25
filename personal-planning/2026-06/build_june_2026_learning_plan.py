# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = OUTPUT_DIR / "june-2026-learning-practice-plan.docx"
REPORT_DATE = "2026-05-25"

FONT_BODY = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 45, 60)
MUTED = RGBColor(92, 104, 116)
GRAY_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF4D6"
GREEN_FILL = "EAF4EA"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text="", style=None, size=11, bold=False, color=INK, after=6, line=1.25):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run_font(r, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    p = add_para(doc, text, style="List Bullet", after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    return p


def add_number(doc, text):
    p = add_para(doc, text, style="List Number", after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    return p


def add_callout(doc, label, text, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ": ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths, header_fill=GRAY_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            if idx == 0 and len(str(value)) <= 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, color=INK, bold=(idx == 0))
    doc.add_paragraph()
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Title", 22, BLUE),
        ("Subtitle", 12, MUTED),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("June 2026 Learning Practice Plan")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Personal Planning | {REPORT_DATE}")
    set_run_font(r, size=9, color=MUTED)
    return doc


def add_cover(doc):
    p = doc.add_paragraph(style="Title")
    r = p.add_run("2026 年 6 月学习实践计划书")
    set_run_font(r, size=22, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(style="Subtitle")
    r = p.add_run("从学术小白到“视频流 + Agent”最小研究闭环")
    set_run_font(r, size=12, color=MUTED)
    p.paragraph_format.space_after = Pt(12)

    add_table(
        doc,
        ["项目", "内容"],
        [
            ("适用对象", "仅系统上过机器学习、深度学习课程，论文阅读与工程复现经验较少，流媒体主链路技术基础有限。"),
            ("6 月定位", "不是冲刺论文创新，而是建立可持续的学习系统：会读论文、会跑视频处理脚本、会把视频转成可检索事件、会记录证据。"),
            ("主线方向", "长期围绕“视频流输入 -> 抽帧/辅助文本 -> 事件索引 -> Agent 检索问答/诊断”推进。"),
            ("保底产出", "2 篇论文精读笔记、1 个 FFmpeg/OpenCV 抽帧实验、1 份事件 schema、1 个可检索视频事件表。"),
            ("扩展产出", "完成 Topic A 最小原型：输入一段视频，生成事件记录，支持简单查询并返回时间戳/关键帧证据。"),
        ],
        [1900, 7460],
        BLUE_FILL,
    )
    add_callout(
        doc,
        "总判断",
        "6 月不要把目标设成“成为 CV 研究者”或“读懂所有前沿论文”。更现实也更有价值的目标，是把论文阅读、视频工程和 Agent 实验接成一个小闭环。只要这个闭环能跑通，你后面 7-9 月的开题、复现和导师沟通都会稳很多。",
        CAUTION_FILL,
    )


def add_current_state(doc):
    add_heading(doc, "1. 你现在的起点：不要把自己误判成零基础", 1)
    add_para(
        doc,
        "你不是完全零基础。你有计算机本科背景、嵌入式/C++经验、视频插件工作经验，也上过机器学习和深度学习课程。真正缺的是三件事：第一，论文阅读的路径感；第二，视频流媒体主链路的可操作经验；第三，把 AI/CV/Agent 工具组合成实验系统的实践经验。",
    )
    add_table(
        doc,
        ["维度", "当前状态", "6 月应对方式"],
        [
            ("论文", "能理解课程概念，但缺少“为什么读、读到什么程度、如何转成实验”的经验。", "采用问题驱动读法：每篇论文只抓问题、方法、数据、指标、可复现模块、对自己题目的启发。"),
            ("AI/CV", "有 ML/DL 课程基础，但模型训练、数据处理、评估脚本经验不足。", "先不训练大模型，优先用现成模型/API/OpenCV/轻量检测工具做视频事件化。"),
            ("流媒体", "当前接触视频插件，但对 FFmpeg、RTSP/RTMP、抽帧、转码、时间戳、码率等主链路仍需补齐。", "用 FFmpeg 命令和 Python/OpenCV 做最小可运行实验，先会观察和记录，再谈架构。"),
            ("Agent", "有兴趣，但容易停留在聊天机器人或框架名词。", "Agent 必须调用真实工具：search_events、get_frame、summarize_clip、verify_answer。"),
            ("时间", "工作日约 2 小时，周末每天约 6 小时。", "每周只设一个主任务，产出优先于收藏资料。"),
        ],
        [1400, 3600, 4360],
    )
    add_callout(
        doc,
        "心理锚点",
        "你的 6 月不是补课月，而是建立“能持续产出”的实践月。学术小白最怕一开始就读太难、做太大、拖太久；所以本计划把论文、技术和项目都压成小步快跑。",
        GREEN_FILL,
    )


def add_paper_strategy(doc):
    add_heading(doc, "2. 论文阅读：基础论文和前沿论文怎么取舍", 1)
    add_para(
        doc,
        "建议采用“夹心读法”，不是从最基础一路顺序读到前沿，也不是直接硬啃最前沿论文。你需要先读 1-2 篇能建立地图的基础/综述/benchmark，再读 2-3 篇与你开题高度相关的前沿方法论文，最后把论文拆成可实现模块。",
    )
    add_table(
        doc,
        ["层级", "6 月读什么", "读到什么程度", "不做什么"],
        [
            ("入门地图", "长视频理解、Video-RAG 或多模态 RAG 的概览/benchmark，如 LongVideoBench、Video-MME。", "知道问题为什么难：长上下文、证据定位、成本、时序结构、评估指标。", "不追求看懂所有模型细节和排行榜。"),
            ("方法主线", "Video-RAG、DrVideo 这类把视频转成文档/辅助文本/证据检索的论文。", "抓住可复现思想：抽帧、辅助文本、索引、检索、证据返回。", "不复现完整 benchmark，不训练大模型。"),
            ("Agent 加层", "HM-RAG、RAVEN、LVAgent 或相关 Agentic Video/RAG 工作。", "只看 Agent 如何拆任务、调工具、合并证据、做验证。", "不陷入多 Agent 框架堆砌。"),
            ("工程补充", "FFmpeg/OpenCV/向量检索/Agents SDK 或 LangGraph 文档。", "能转成实验脚本和工具函数。", "不把文档学习当作独立目标。"),
        ],
        [1500, 3300, 3100, 1460],
    )

    add_heading(doc, "2.1 6 月论文清单", 2)
    add_table(
        doc,
        ["优先级", "论文/资料", "本月任务", "输出"],
        [
            ("必读 1", "LongVideoBench 或 Video-MME", "了解长视频/多模态视频评估怎么设计问题、指标和数据。", "1 页中文笔记：问题、指标、可借鉴处。"),
            ("必读 2", "Video-RAG", "理解低成本长视频理解的基本管线：辅助文本、检索增强、回答。", "1 页精读笔记 + 画出自己的轻量版管线图。"),
            ("必读 3", "DrVideo", "学习“视频转文档/证据检索”的思路。", "提炼可复现模块：video document、retrieval、evidence。"),
            ("选读 1", "HM-RAG 或 RAVEN", "只看 Agent/图结构/层级记忆如何组织视频证据。", "列出 3 个可以放进自己系统的工具函数。"),
            ("工程文档", "FFmpeg 官方文档、OpenCV VideoCapture、向量库入门文档。", "边查边做，不要求系统读完。", "命令记录 + 可运行脚本。"),
        ],
        [1100, 2600, 3700, 1960],
        BLUE_FILL,
    )

    add_heading(doc, "2.2 每篇论文固定读法", 2)
    steps = [
        "第一遍 20 分钟：只读摘要、引言、图 1、结论，写下这篇论文解决什么痛点。",
        "第二遍 60-90 分钟：读方法和实验，标出输入、输出、模块、指标、数据集。",
        "第三遍 30 分钟：只回答一个问题：它能给我的 Topic A/D 原型带来什么可实现模块？",
        "最后 20 分钟：写 6 行笔记，不写长篇翻译。每篇论文必须产出一张“方法拆解表”。",
    ]
    for step in steps:
        add_number(doc, step)
    add_callout(
        doc,
        "阅读边界",
        "6 月读论文的目标不是证明你学术能力强，而是帮你做出最小原型。凡是不能转成问题定义、系统模块、实验指标或导师沟通材料的内容，本月都可以先跳过。",
        CAUTION_FILL,
    )


def add_technical_plan(doc):
    add_heading(doc, "3. 技术基础学习与实践", 1)
    add_para(
        doc,
        "技术学习不要按课程目录推进，而要按系统链路推进。6 月只打通一个最小链路：拿到视频 -> 抽帧/截图片段 -> 生成结构化事件 -> 建索引 -> 查询返回证据。C++ 可以保留为长期护城河，但 6 月原型优先用 Python，因为它更适合快速实验和论文复现。",
    )
    add_table(
        doc,
        ["模块", "本月学习内容", "实践任务", "验收标准"],
        [
            ("环境", "Python、uv/venv、Git、README、实验目录结构。", "建立 video-agent-lab 项目目录，记录依赖和运行命令。", "别人按 README 能跑通一个脚本。"),
            ("视频输入", "FFmpeg 命令、容器/编码、帧率、时间戳、RTSP/本地文件。", "用 FFmpeg 对 1-2 段视频做信息查看、抽帧、截取片段。", "能解释 fps、duration、bitrate、pts/dts 的基本含义。"),
            ("抽帧与证据", "OpenCV VideoCapture、关键帧/间隔抽帧、截图命名规范。", "每隔 N 秒抽 1 帧，保存 frame_id、timestamp、path。", "生成 frames/ 与 metadata.csv/json。"),
            ("事件 schema", "把视觉/语音/文本结果统一成事件记录。", "设计 event_id、start_time、end_time、objects、description、evidence_frames、source。", "能把一段视频转成 10-30 条人工/半自动事件。"),
            ("检索", "SQLite/JSONL、关键词检索、向量检索的最小概念。", "先做关键词检索，再扩展 embedding 检索。", "输入问题，返回 Top-K 事件和时间戳。"),
            ("Agent", "工具调用、结构化输出、证据验证。", "写 3 个工具函数：search_events、get_evidence、answer_with_citations。", "回答必须带 timestamp 和 evidence_frame。"),
        ],
        [1300, 2700, 3200, 2160],
    )

    add_heading(doc, "3.1 推荐项目结构", 2)
    add_para(doc, "建议 6 月新建一个独立实验仓库或子目录，结构保持简单：")
    for item in [
        "data/raw/：原始视频，优先放公开或自采无隐私视频。",
        "data/frames/：抽帧结果，文件名包含 video_id 和 timestamp。",
        "data/events.jsonl：事件记录，一行一个事件。",
        "src/extract_frames.py：抽帧脚本。",
        "src/build_events.py：人工/半自动事件生成脚本。",
        "src/search_events.py：检索脚本。",
        "notes/papers/：论文笔记。",
        "README.md：本周能跑什么、怎么跑、还有什么问题。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.2 最小事件 schema", 2)
    add_table(
        doc,
        ["字段", "含义", "6 月是否必须"],
        [
            ("video_id", "视频编号或文件名。", "必须"),
            ("event_id", "事件编号。", "必须"),
            ("start_time/end_time", "事件时间范围，单位秒。", "必须"),
            ("description", "事件自然语言描述。", "必须"),
            ("evidence_frames", "关键帧路径列表。", "必须"),
            ("objects/entities", "人物、物体、地点、设备等实体。", "建议"),
            ("source", "人工标注、OCR、ASR、检测模型或规则。", "必须"),
            ("confidence", "置信度或主观可靠性说明。", "建议"),
        ],
        [1800, 5200, 2360],
    )


def add_june_goals(doc):
    add_heading(doc, "4. 6 月目标：保底目标与扩展目标", 1)
    add_table(
        doc,
        ["类别", "保底目标", "扩展目标"],
        [
            ("论文", "精读 2 篇：Video-RAG + LongVideoBench/Video-MME 二选一；每篇 1 页笔记。", "精读 4 篇：再加入 DrVideo 和 HM-RAG/RAVEN；形成 related work 小地图。"),
            ("视频工程", "跑通 FFmpeg 信息查看、抽帧、截取片段；整理 10 条命令记录。", "增加 RTSP/本地流模拟，记录延迟、帧率、丢帧等观察。"),
            ("数据化", "为 1 段 3-10 分钟视频手工/半自动生成 10-30 条事件。", "为 2-3 段视频生成事件，并统一 schema。"),
            ("检索", "用关键词检索返回相关事件和时间戳。", "加入向量检索或混合检索，输出 Top-K 证据。"),
            ("Agent", "先不用复杂框架，写一个函数式 answer_with_evidence。", "用 Agents SDK/LangGraph 做工具调用 Demo。"),
            ("导师沟通", "月底整理 1 页 Topic A/D 验证结论。", "带着可运行 Demo 和方法图与导师沟通。"),
        ],
        [1300, 4000, 4060],
        BLUE_FILL,
    )
    add_callout(
        doc,
        "6 月成功标准",
        "月底你能说清楚三件事就算成功：我读过哪些论文，它们给我的系统带来什么模块；我能如何把视频变成结构化事件；我下一步更适合走 Topic A 还是 Topic D。",
        GREEN_FILL,
    )


def add_weekly_plan(doc):
    add_heading(doc, "5. 四周执行计划", 1)
    add_table(
        doc,
        ["周次", "主题", "工作日任务", "周末任务", "交付物"],
        [
            ("第 1 周", "建立地图与环境", "读 LongVideoBench/Video-MME；安装工具；收集 1-2 段公开视频。", "跑通 FFmpeg 信息查看、抽帧、截图；写实验 README。", "论文笔记 1；FFmpeg 命令记录；frames/ 示例。"),
            ("第 2 周", "Video-RAG 思想转工程", "精读 Video-RAG；设计自己的轻量管线图。", "写 extract_frames.py 和 events.jsonl 草稿；手工标 10 条事件。", "论文笔记 2；事件 schema v0；10 条事件。"),
            ("第 3 周", "检索与证据返回", "读 DrVideo；实现关键词检索；整理时间戳证据。", "写 search_events.py；准备 5 个问题测试。", "Top-K 检索结果；问题-证据表；DrVideo 笔记。"),
            ("第 4 周", "Agent/D 方向小验证", "选读 HM-RAG/RAVEN 或日志诊断资料；写工具函数接口。", "实现 answer_with_evidence；整理 Topic A/D 对比结论。", "最小问答 Demo；月底复盘；7 月论文清单。"),
        ],
        [1000, 1600, 2900, 2900, 960],
    )

    add_heading(doc, "5.1 每周固定节奏", 2)
    add_table(
        doc,
        ["时间", "建议安排", "产出要求"],
        [
            ("周一 2h", "复盘上周，确定本周唯一主任务。", "写 5 行周计划。"),
            ("周二 2h", "视频工程实验：FFmpeg/OpenCV/脚本。", "提交或保存一个可运行变化。"),
            ("周三 2h", "论文阅读。", "更新论文拆解表。"),
            ("周四 2h", "检索/Agent/数据结构实验。", "新增一个函数或一个样例。"),
            ("周五 2h", "整理笔记、命令、问题清单。", "README 或学习日志更新。"),
            ("周六 6h", "长项目块：把本周主任务跑通。", "可运行脚本或 Demo。"),
            ("周日 6h", "上午补缺口，下午写文档和复盘，晚上休息。", "周复盘 + 下周待办。"),
        ],
        [1300, 5400, 2660],
    )


def add_risk_controls(doc):
    add_heading(doc, "6. 风险控制与降级规则", 1)
    add_table(
        doc,
        ["风险", "表现", "降级处理"],
        [
            ("读论文卡住", "公式、模型细节、实验设置看不懂，读 3 小时没有输出。", "只读摘要/图/方法框架，先写“我能复用什么”，跳过细节。"),
            ("技术范围爆炸", "同时想学 FFmpeg、GStreamer、WebRTC、VLM、Agent 框架。", "6 月只保留 FFmpeg/OpenCV/JSONL/关键词检索。其他放入 later list。"),
            ("Demo 做不出来", "抽帧、索引、查询任一环节阻塞超过 2 天。", "改成人工事件表 + 简单检索，先保证闭环。"),
            ("过度追前沿", "论文越看越新，但代码没有进展。", "每读 1 篇论文必须转成 1 个字段、1 个模块或 1 个实验问题。"),
            ("自我否定", "觉得自己学术太弱、基础太差，不敢开始。", "把任务缩到 30 分钟：跑一条命令、读一张图、写 5 行事件。"),
        ],
        [1600, 3800, 3960],
    )
    add_callout(
        doc,
        "关键提醒",
        "6 月任何时候都可以降级，但不要中断。保底路线跑通，比完美路线停在收藏夹里强得多。",
        CAUTION_FILL,
    )


def add_checklists(doc):
    add_heading(doc, "7. 月底检查表", 1)
    checks = [
        "[ ] 我是否完成至少 2 篇论文笔记？",
        "[ ] 我是否能用自己的话解释 Video-RAG/DrVideo 为什么适合我的题目？",
        "[ ] 我是否跑通过 FFmpeg 抽帧、截取片段、查看视频信息？",
        "[ ] 我是否为至少 1 段视频生成了 events.jsonl 或等价事件表？",
        "[ ] 我是否能输入一个问题，返回相关事件、时间戳和关键帧证据？",
        "[ ] 我是否整理了 Topic A 与 Topic D 的适配度对比？",
        "[ ] 我是否写下 7 月要读的 6-8 篇论文清单？",
    ]
    for check in checks:
        add_bullet(doc, check)

    add_heading(doc, "8. 参考论文与资料入口", 1)
    add_table(
        doc,
        ["资料", "用途"],
        [
            ("LongVideoBench / Video-MME", "建立长视频理解与多模态视频评估的地图。"),
            ("Video-RAG", "学习低成本长视频检索增强思路，是 6 月最重要的方法论文。"),
            ("DrVideo", "学习把长视频转成文档和证据检索的问题设定。"),
            ("HM-RAG / RAVEN", "观察 Agent、层级记忆、结构化视频表示如何进入视频理解系统。"),
            ("FFmpeg 官方文档", "视频输入、转码、抽帧、截取片段、流处理基础。"),
            ("OpenCV VideoCapture", "快速读取视频帧和构建 Python 实验脚本。"),
            ("OpenAI Agents SDK / LangGraph", "后续做工具调用与 Agent 工作流时参考。"),
        ],
        [2600, 6760],
    )


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_doc()
    add_cover(doc)
    add_current_state(doc)
    add_paper_strategy(doc)
    add_technical_plan(doc)
    add_june_goals(doc)
    add_weekly_plan(doc)
    add_risk_controls(doc)
    add_checklists(doc)
    doc.core_properties.title = "2026 年 6 月学习实践计划书"
    doc.core_properties.subject = "Video stream plus agent learning and practice plan"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    build()
