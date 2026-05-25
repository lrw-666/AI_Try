from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = OUTPUT_DIR / "long-video-agent-thesis-feasibility-report.docx"
REPORT_DATE = "2026-05-25"

FONT_BODY = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 45, 60)
MUTED = RGBColor(92, 104, 116)
GRAY_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF4D6"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text="", style=None, size=11, bold=False, color=INK, after=6, line=1.25):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run_font(r, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    p = add_para(doc, text, style="List Bullet", after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    return p


def add_callout(doc, label, text, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ": ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths, header_fill=GRAY_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            if idx == 0 and len(str(value)) <= 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, color=INK, bold=(idx == 0))
    doc.add_paragraph()


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Title", 22, BLUE),
        ("Subtitle", 12, MUTED),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Long Video Agent Thesis Feasibility")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Thesis Feasibility Report | {REPORT_DATE}")
    set_run_font(r, size=9, color=MUTED)
    return doc


def add_cover(doc):
    p = doc.add_paragraph(style="Title")
    r = p.add_run("Long Video Agent Thesis Feasibility Report")
    set_run_font(r, size=22, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(style="Subtitle")
    r = p.add_run("面向长时视频流的语义事件索引与 Agent 检索问答系统：就业、AI 冲击与论文前沿分析")
    set_run_font(r, size=12, color=MUTED)
    p.paragraph_format.space_after = Pt(12)

    add_table(
        doc,
        ["项目", "内容"],
        [
            ("分析对象", "“面向长时视频流的语义事件索引与 Agent 检索问答系统”作为硕士论文与职业作品主线的可行性。"),
            ("个人定位", "C++ 视频系统 + AI Agent 工程化；用视频流作为通往 XR、机器人、智慧城市和多模态应用的底层入口。"),
            ("核心问题", "毕业后是否匹配市场岗位、是否抗 AI 冲击、论文是否有前沿性与可创新空间、9 月开题是否可落地。"),
            ("结论等级", "建议作为 2026-2028 年论文主线的优先方向，但必须把题目收束为“系统方法 + 可量化评估”，避免泛泛做聊天机器人。"),
            ("生成日期", REPORT_DATE),
        ],
        [1900, 7460],
        BLUE_FILL,
    )
    add_callout(
        doc,
        "一句话结论",
        "这个方向对你是“职业与论文可以共用的一条路”：它能从当前视频插件工作自然生长出来，又能接上 Agent、RAG、多模态视频理解和成都 AI 应用岗位；但论文创新不能写成“把现有模型串起来”，必须围绕长视频事件索引、证据定位、Agent 工具调用和评估闭环形成明确贡献。",
        CAUTION_FILL,
    )


def add_market_feasibility(doc):
    add_heading(doc, "1. 就业可行性判断", 1)
    add_para(
        doc,
        "这个方向对应的岗位不是单一岗位名，而是一组正在融合的岗位族：视频 AI 工程师、AI Agent 工程师、多模态应用工程师、流媒体/视频平台工程师、AI 应用后端工程师。你的优势在于 C++ 视频系统和 Agent 工程的交叉，而不是单纯去和算法博士竞争视频大模型训练。",
    )
    add_table(
        doc,
        ["岗位族", "市场常见要求", "你的匹配度", "需要补齐"],
        [
            ("视频 AI 工程师", "C/C++ 或 Python；FFmpeg/GStreamer；视频接入、抽帧、转码、边缘/云端推理；OpenCV/ONNX/PyTorch；性能与稳定性。", "中高。你有 C++ 和视频插件基础，但需要把播放边缘业务扩展为流处理、抽帧、推拉流和 AI 推理链路。", "FFmpeg/GStreamer、RTSP/RTMP/WebRTC、ONNX Runtime、视频数据 schema。"),
            ("AI Agent 工程师", "Python/TypeScript；RAG；LangGraph/Agents SDK；工具调用；向量数据库；后端服务；评估、监控、权限和错误恢复。", "中等偏高。你有系统思维和工程敏感度，但需要补 AI 应用后端和 Agent 生产化。", "FastAPI、LangGraph/OpenAI Agents SDK、Milvus/FAISS、Agent tracing/evaluation。"),
            ("多模态应用工程师", "视觉/语音/文本多模态模型；视频问答；OCR/ASR；检索增强；数据处理；模型调用与评测。", "中等。导师视觉方向能帮你建立论文入口，但要避免走纯模型训练路线。", "视频理解基准、Video-RAG 方法、OCR/ASR/检测标签融合。"),
            ("视频平台工程师", "流媒体协议、媒体服务器、低延迟、编解码、跨平台、稳定性、监控与故障定位。", "中高。这是你的保底工程壁垒。", "更系统的音视频主链路经验，不只做业务插件。"),
            ("技术产品/解决方案", "能把 AI 视频检索、事件分析、智能运维、数字孪生或机器人视觉流讲成场景方案。", "中等。可作为后续转型，不建议现在完全放弃技术。", "PRD、方案表达、行业案例、非强销售型沟通能力。"),
        ],
        [1700, 3300, 2360, 2000],
    )
    add_callout(
        doc,
        "毕业后岗位判断",
        "如果到 2028 年你只有“论文 + 简单 Demo”，市场竞争力一般；如果你能展示一套可运行系统：视频流接入、事件抽取、向量/时间索引、Agent 查询、证据帧返回、延迟与准确率评估，那么它能同时服务 AI Agent、视频 AI 和流媒体工程岗位。",
    )

    add_heading(doc, "2. 成都与城市选择", 1)
    add_para(
        doc,
        "成都不是这个方向的最强岗位城市，但并非没有落地基础。四川省政府公开资料显示，成都 2025 年提出人工智能核心产业规模 1300 亿元目标，并强调大模型、机器人、场景应用、数据集、典型应用场景和产业园建设；成都元宇宙行动方案也涉及人工智能、感知交互、数字孪生等方向。这说明本方向可挂靠成都的 AI 应用、智慧城市、机器人、数字孪生、文旅科技和企业数字化场景。",
    )
    add_bullet(doc, "成都优势：生活性价比、AI/机器人/数字孪生政策场景、政企数字化、智慧城市和应用型企业。")
    add_bullet(doc, "成都风险：高质量多模态视频 Agent 岗位密度不如北京、上海、深圳、杭州；可能更多是项目交付或传统场景。")
    add_bullet(doc, "建议策略：2027 年下半年开始做岗位地图；成都优先，但保留一线跳板和远程岗位，避免毕业时被城市单点限制。")


def add_ai_impact(doc):
    add_heading(doc, "3. AI 冲击下的发展形态", 1)
    add_para(
        doc,
        "AI 对这条路线的影响不是“替代”那么简单。普通视频播放、简单代码实现、基础 RAG Demo 会被 AI 快速压缩；但视频流系统、数据结构化、证据定位、可靠 Agent 工作流、评估体系和业务落地会变得更重要。",
    )
    add_table(
        doc,
        ["能力类型", "被 AI 冲击程度", "原因", "你的应对"],
        [
            ("普通编码", "高", "代码生成工具会降低实现门槛，简单 CRUD、脚手架和 API 调用不再稀缺。", "不要把自己定位成只会写接口的人。"),
            ("视频系统工程", "中低", "视频流、时间戳、延迟、丢帧、编解码、协议兼容和性能调优仍有强工程复杂度。", "把 C++/FFmpeg/GStreamer/流媒体作为工程底座。"),
            ("Agent Demo", "高", "Prompt + RAG + UI 的浅层项目会严重同质化。", "必须做工具调用、状态、评估、错误恢复、证据可追溯。"),
            ("视频语义索引", "中", "模型能自动生成描述，但长视频事件切片、证据定位和索引结构仍需设计。", "围绕“可检索、可解释、可评估”做论文创新。"),
            ("系统评估与落地", "低", "业务指标、失败模式、数据闭环、成本控制和用户场景需要工程判断。", "把论文实验和求职作品都做成可量化系统。"),
        ],
        [1700, 1600, 3100, 2960],
    )
    add_callout(
        doc,
        "抗冲击定位",
        "未来更有价值的不是“会调用大模型分析视频”，而是“能把长视频变成可索引、可检索、可问答、可追溯证据、可评估质量的系统”。这正好避开纯算法重数学，也避开低端监控运维。",
        BLUE_FILL,
    )


def add_research_feasibility(doc):
    add_heading(doc, "4. 论文创新性与前沿性", 1)
    add_para(
        doc,
        "这个方向有前沿性。2024-2026 年，长视频理解、Video-RAG、视频时间定位、多模态 Agent、视频库问答和长视频基准快速增长。它适合非全硕士做应用型研究，因为你可以不训练大模型，而是设计一个更工程可落地、更可评估的系统方法。",
    )
    add_table(
        doc,
        ["可创新点", "为什么不是简单拼接", "适合你的实现方式"],
        [
            ("多源语义事件索引", "长视频不是一段文本，事件需要由帧、OCR、ASR、目标、场景、时间戳共同组成。", "设计事件 schema：时间段、证据帧、文本描述、对象、声音、置信度、来源。"),
            ("时间索引 + 向量索引混合检索", "只做向量检索容易丢失时序；只做时间切片又无法语义查询。", "使用时间窗口、关键词、向量 embedding、事件类型共同召回。"),
            ("Agent 工具调用闭环", "Agent 不直接“看完视频回答”，而是调用检索、定位、证据汇总、报告生成工具。", "实现可观测工具链，并记录每次检索和证据来源。"),
            ("证据可追溯问答", "视频问答不能只给自然语言答案，还要返回帧、时间段和依据。", "答案中附带 timestamp、key frame、事件片段和置信说明。"),
            ("轻量化与低成本", "前沿大模型能处理长上下文，但成本高；工程场景需要低成本索引和按需检索。", "比较全视频输入、抽帧输入、事件索引输入三种成本和效果。"),
            ("面向特定场景的评估", "公开基准未必覆盖你的业务场景，应用型论文可以建立小规模可控数据集。", "自建 3-5 小时视频样例，标注事件和问题，做 Recall@K、定位误差、问答正确率。"),
        ],
        [2100, 3600, 3660],
    )

    add_heading(doc, "5. 建议收束后的论文题目", 1)
    add_callout(
        doc,
        "推荐题目",
        "面向长时视频流的语义事件索引与 Agent 检索问答系统研究。副标题可写为：基于多源辅助文本与混合检索的长视频证据定位方法。",
        CAUTION_FILL,
    )
    add_table(
        doc,
        ["论文模块", "建议写法"],
        [
            ("研究背景", "长视频信息量大，直接输入多模态大模型成本高且难定位证据；工业、城市、教育、会议、机器人等场景需要按语义查询视频事件。"),
            ("研究问题", "如何把长时视频流转化为可检索事件；如何让 Agent 按问题调用工具检索证据；如何评价答案是否准确且证据可追溯。"),
            ("方法设计", "视频切片与抽帧、多源辅助文本生成、语义事件 schema、混合检索、Agent 工具链、证据生成与报告。"),
            ("实验设计", "公开数据 + 自建小数据；与 naive RAG、直接帧输入、仅 ASR/OCR/检测标签方法对比；做消融实验。"),
            ("主要指标", "Recall@K、时间定位误差、证据帧命中率、问答准确率、平均响应时间、索引成本、人工评分。"),
            ("预期贡献", "提出一个低成本长视频语义事件索引方法；实现可追溯 Agent 问答系统；在小规模场景验证有效性。"),
        ],
        [1800, 7560],
    )


def add_paper_review(doc):
    add_heading(doc, "6. 重点论文与阅读建议", 1)
    add_para(
        doc,
        "严格说，这个方向的大量关键工作还在顶会和预印本阶段；如果只限定“顶刊”，会错过前沿。建议按“顶刊/综述 + 顶会/基准 + 前沿系统论文”三层阅读。下面列出的论文均与开题相关，可作为 6-9 月论文笔记主干。",
    )
    add_table(
        doc,
        ["层级", "论文", "主要内容", "与你题目的关系"],
        [
            ("顶刊综述", "Wu et al., A Survey on Video Temporal Grounding With Multimodal Large Language Model, IEEE TPAMI, 2026.", "系统梳理多模态大模型驱动的视频时间定位，讨论 MLLM 在时间理解、训练范式、视频特征处理、基准和未来方向上的作用。", "给你的论文提供“长视频证据定位/时间定位”的学术背景，可解释为什么答案必须带时间戳和证据。"),
            ("顶刊研究", "HM-RAG: Long Video Reasoning and Anomaly Detection via Hierarchical Multi-Agent RAG, Pattern Recognition, 2026.", "将层次化多 Agent 与 RAG 结合，用时间记忆 Agent、外部知识检索和知识融合处理长视频推理与异常检测。", "与你的 Agent 检索问答系统非常接近，可作为最重要的对标论文之一；但你可以做更轻量、更工程化的事件索引。"),
            ("顶会方法", "Ma et al., DrVideo: Document Retrieval Based Long Video Understanding, CVPR 2025.", "将长视频理解转化为长文档检索问题，先把视频变成文本型文档，再检索关键帧并用增强信息更新文档。", "证明“视频转文档/事件索引”是顶会认可方向；你的系统可从文档检索扩展到 Agent 工具链和证据问答。"),
            ("顶会基准", "LongVideoBench: A Benchmark for Long-context Interleaved Video-Language Understanding, NeurIPS 2024.", "构建最长约一小时的视频语言交错问答基准，强调长上下文视频理解和引用推理问题。", "可作为评估设计参考，说明长视频不是短视频理解的简单放大。"),
            ("顶会基准", "Video-MME: Comprehensive Evaluation Benchmark of MLLMs in Video Analysis, CVPR 2025.", "覆盖多领域、多时长、多任务的视频理解评测，用于评估多模态大模型在视频分析中的能力。", "可作为开题中引用的通用视频理解基准，帮助你设计任务维度和指标。"),
            ("前沿方法", "Video-RAG: Visually-aligned Retrieval-Augmented Long Video Comprehension, arXiv 2024.", "用音频、OCR、目标检测等辅助文本增强视觉对齐，训练成本低，可插拔到现有 LVLM，并在多个长视频基准上提升效果。", "与你的“多源辅助文本 + 语义事件索引”高度相关，是最直接的技术参考。"),
            ("前沿系统", "VideoRAG: Retrieval-Augmented Generation with Extreme Long-Context Videos, arXiv 2025.", "提出双通道架构：图式文本知识 grounding + 多模态上下文编码，并构建跨视频知识图以处理超长视频。", "适合作为上限参考；你不必实现它全部复杂度，但可借鉴图结构和跨视频关系。"),
            ("视频库问答", "Towards Retrieval Augmented Generation over Large Video Libraries, arXiv 2024.", "提出 Video Library QA，使用 LLM 生成搜索查询，检索视频时刻，再生成带时间戳的回答。", "非常贴近“视频库检索问答”应用，适合用来说明职业场景。"),
            ("Agent 方向", "RAVEN: An Agentic Framework for Multimodal Entity Discovery from Large-Scale Video Collections, arXiv 2025.", "用 Agent 进行视频集合中的实体发现、schema 生成和结构化表示，面向大规模视频集合检索。", "说明 Agent 不只是问答 UI，而可以负责 schema、抽取和结构化知识。"),
            ("机器人场景", "Multi-RAG: A Multimodal RAG System for Adaptive Video Understanding, arXiv 2025.", "整合视频、音频、文本等多源信息，服务动态场景中的人机辅助理解。", "可作为未来向机器人/具身智能扩展的桥梁，但不必现在深做机器人底层。"),
        ],
        [1200, 2800, 3000, 2360],
    )
    add_callout(
        doc,
        "阅读顺序建议",
        "先读 DrVideo、Video-RAG、LongVideoBench 和 TPAMI Survey；再读 HM-RAG、VideoRAG、RAVEN。这样你先建立“长视频为何难、如何评估、如何索引”的主线，再看多 Agent 和图检索等高级结构。",
    )


def add_opening_plan(doc):
    add_heading(doc, "7. 9 月开题可行性", 1)
    add_table(
        doc,
        ["时间", "必须完成", "开题价值"],
        [
            ("2026.06", "完成 8 篇论文初读：DrVideo、Video-RAG、LongVideoBench、Video-MME、TPAMI Survey、HM-RAG、VideoRAG、VLQA。", "能证明方向前沿且不是拍脑袋。"),
            ("2026.07", "跑通最小系统：视频输入 -> 抽帧 -> OCR/ASR/检测标签之一 -> 事件表 -> 文本检索。", "证明工程可行。"),
            ("2026.08", "加入向量检索和 Agent 工具调用，回答 10 个手工问题并返回时间戳/证据帧。", "证明论文方法可演示。"),
            ("2026.09", "形成开题报告：研究问题、相关工作、系统架构、数据、指标、风险、计划。", "让导师看到这是应用型研究而不是普通项目。"),
        ],
        [1500, 5400, 2460],
    )

    add_heading(doc, "8. 最小可行实验设计", 1)
    add_table(
        doc,
        ["模块", "最低版本", "进阶版本"],
        [
            ("数据", "3-5 小时公开视频或自采无隐私视频，手工标注 50-100 个事件和 50 个问题。", "加入更长视频或多视频集合，扩展到 200+ 问题。"),
            ("事件抽取", "抽帧 + ASR/OCR/检测标签中的 2 类辅助文本。", "加入场景分类、运动事件、目标跟踪或关键帧聚类。"),
            ("检索", "关键词检索 + 向量检索 + 时间窗口过滤。", "混合召回排序、图关系、查询改写。"),
            ("Agent", "工具调用：search_events、get_frame、summarize_evidence、answer_with_citations。", "多 Agent：检索 Agent、证据 Agent、回答 Agent、评估 Agent。"),
            ("指标", "Recall@K、时间定位误差、问答准确率、延迟。", "加证据质量人工评分、成本、消融实验。"),
        ],
        [1500, 3900, 3960],
    )


def add_personal_recommendation(doc):
    add_heading(doc, "9. 对你的个人建议", 1)
    add_bullet(doc, "不要把论文做成“调用某个大模型分析视频”的浅层应用。那样就业和答辩都不够稳。")
    add_bullet(doc, "要把核心能力放在视频流工程、事件索引、检索结构、Agent 工具链和评估闭环上。")
    add_bullet(doc, "C++ 不是丢掉，而是作为视频底座；Python 是 AI/Agent/论文实验加速器。")
    add_bullet(doc, "工作中即使只接触边缘视频业务，也要主动抽象视频链路、日志、数据、插件架构和可扩展点。")
    add_bullet(doc, "论文选题要和导师沟通“视觉方向 + 应用系统 + 可量化实验”，不要只说 Agent，否则容易显得偏软件工程。")
    add_bullet(doc, "毕业后投递时，把自己包装成“懂视频系统的 AI Agent 工程师”，而不是泛 AI 应用开发或普通 C++ 业务开发。")

    add_callout(
        doc,
        "最终判断",
        "该方向值得进入 6-9 月开题验证。它不是最容易的路，但它是目前最能同时服务你兴趣、论文、工作经验、成都就业和 AI 时代抗冲击的一条复合路线。",
        BLUE_FILL,
    )


def add_sources(doc):
    add_heading(doc, "10. 主要资料来源", 1)
    sources = [
        "OpenAI Agents SDK 文档：Agent 可使用工具、handoff、流式输出、状态和 trace 组织工作流。https://developers.openai.com/api/docs/guides/agents",
        "LangGraph workflows and agents 文档：强调工具调用、结构化输出、记忆、持久化、流式处理、调试和部署。https://docs.langchain.com/oss/python/langgraph/workflows-agents",
        "FFmpeg 官方文档：FFmpeg 可处理文件、网络流和采集设备输入，是视频流处理基础工具。https://www.ffmpeg.org/ffmpeg.html",
        "ONVIF Profile T：覆盖 IP 视频系统中的视频编码、成像设置、事件、运动告警、元数据流和 PTZ 等能力。https://www.onvif.org/profiles/profile-t/",
        "BLS Occupational Outlook Handbook：软件开发、质量保障和测试岗位 2024-2034 年预计增长 15%，但传统 programmer 岗位趋势较弱。https://www.bls.gov/ooh/computer-and-information-technology/software-developers.htm",
        "GitHub Octoverse 2024：AI 推动 Python 和相关开源活动增长，说明 AI 原生开发正在成为软件工程生态的一部分。https://github.blog/news-insights/octoverse/octoverse-2024/",
        "四川省人民政府：成都 2025 年人工智能核心产业规模目标 1300 亿元，并强调机器人、行业大模型、场景应用和产业园建设。https://www.sc.gov.cn/10462/10464/10465/10595/2025/3/28/5a9be7787403485083e17b95edae09b0.shtml",
        "四川省人民政府：成都元宇宙行动方案提出数字孪生、感知交互、人工智能等基础核心技术方向。https://www.sc.gov.cn/10462/10464/10465/10595/2023/1/4/fdcc2f41865a4202a332d2903b500098.shtml",
        "Wu et al., A Survey on Video Temporal Grounding With Multimodal Large Language Model, IEEE TPAMI, 2026. https://doi.org/10.1109/TPAMI.2025.3615586",
        "HM-RAG: Long Video Reasoning and Anomaly Detection via Hierarchical Multi-Agent RAG, Pattern Recognition, 2026. https://doi.org/10.1016/j.patcog.2026.113622",
        "DrVideo: Document Retrieval Based Long Video Understanding, CVPR 2025. https://openaccess.thecvf.com/content/CVPR2025/html/Ma_DrVideo_Document_Retrieval_Based_Long_Video_Understanding_CVPR_2025_paper.html",
        "LongVideoBench, NeurIPS 2024. https://proceedings.neurips.cc/paper_files/paper/2024/hash/329ad516cf7a6ac306f29882e9c77558-Abstract-Datasets_and_Benchmarks_Track.html",
        "Video-MME, CVPR 2025 / arXiv 2024. https://arxiv.org/abs/2405.21075",
        "Video-RAG: Visually-aligned Retrieval-Augmented Long Video Comprehension. https://arxiv.org/abs/2411.13093",
        "VideoRAG: Retrieval-Augmented Generation with Extreme Long-Context Videos. https://arxiv.org/abs/2502.01549",
        "Towards Retrieval Augmented Generation over Large Video Libraries. https://arxiv.org/abs/2406.14938",
        "RAVEN: An Agentic Framework for Multimodal Entity Discovery from Large-Scale Video Collections. https://arxiv.org/abs/2504.06272",
        "Multi-RAG: A Multimodal Retrieval-Augmented Generation System for Adaptive Video Understanding. https://arxiv.org/abs/2505.23990",
    ]
    for source in sources:
        add_bullet(doc, source)


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_doc()
    add_cover(doc)
    add_market_feasibility(doc)
    add_ai_impact(doc)
    add_research_feasibility(doc)
    add_paper_review(doc)
    add_opening_plan(doc)
    add_personal_recommendation(doc)
    add_sources(doc)
    doc.core_properties.title = "Long Video Agent Thesis Feasibility Report"
    doc.core_properties.subject = "Feasibility analysis for long video semantic event indexing and agent retrieval QA"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    build()

