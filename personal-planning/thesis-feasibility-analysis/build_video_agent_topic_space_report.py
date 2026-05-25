from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = OUTPUT_DIR / "long-video-agent-thesis-feasibility-report.docx"
REPORT_DATE = "2026-05-25"

FONT_BODY = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 45, 60)
MUTED = RGBColor(92, 104, 116)
GRAY_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF4D6"
GREEN_FILL = "EAF4EA"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text="", style=None, size=11, bold=False, color=INK, after=6, line=1.25):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    p = add_para(doc, text, style="List Bullet", after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    return p


def add_callout(doc, label, text, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ": ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths, header_fill=GRAY_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            if idx == 0 and len(str(value)) <= 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, color=INK, bold=(idx == 0))
    doc.add_paragraph()


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Title", 22, BLUE),
        ("Subtitle", 12, MUTED),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Video Stream + Agent Topic Space")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Thesis Topic Space Report | {REPORT_DATE}")
    set_run_font(r, size=9, color=MUTED)
    return doc


def add_cover(doc):
    p = doc.add_paragraph(style="Title")
    r = p.add_run("Video Stream + Agent Thesis Topic Space Report")
    set_run_font(r, size=22, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(style="Subtitle")
    r = p.add_run("视频流 + Agent 方向前沿、可复现性、创新空间与 5 个论文题目比较")
    set_run_font(r, size=12, color=MUTED)
    p.paragraph_format.space_after = Pt(12)

    add_table(
        doc,
        ["项目", "内容"],
        [
            ("文档目标", "不再只分析单一题目，而是帮助你打开“视频流 + Agent”方向的选题空间，并判断哪些题目适合非全硕士背景。"),
            ("核心约束", "你有 C++/视频插件背景，但视频主链路、AI 后端、Agent 工程和论文实验还需要补；学习时间有限，不能选择重训练大模型或重 SLAM/3D 的路线。"),
            ("主判断", "最适合你的题目应当是系统型、低训练成本、可复现、可量化评估、能连接就业岗位的题目。"),
            ("推荐优先级", "优先考虑：长视频语义事件索引 Agent、在线视频流记忆 Agent、视频工程日志/画面联合诊断 Agent。谨慎考虑需要训练大模型或大规模标注的题目。"),
            ("生成日期", REPORT_DATE),
        ],
        [1900, 7460],
        BLUE_FILL,
    )
    add_callout(
        doc,
        "一句话结论",
        "视频流 + Agent 的前沿已经从“视频问答 Demo”推进到长视频 RAG、多 Agent 协作、实时流式理解、视频库检索和证据可追溯；你可以复现其中的工程化轻量路线，但不适合复现需要大规模训练、GPU 集群或海量标注的路线。你的创新应放在事件结构、检索策略、Agent 工具链、评估闭环和具体场景适配。",
        CAUTION_FILL,
    )


def add_frontier_map(doc):
    add_heading(doc, "1. 视频流 + Agent 前沿研究到了什么地步", 1)
    add_para(
        doc,
        "截至 2026 年 5 月，这个方向大致分成六条前沿线。它们不是互相独立，而是在逐渐融合：长视频理解解决“看得长”，Video-RAG 解决“找得到”，Agent 解决“会调用工具”，在线理解解决“边看边反应”，视频库问答解决“查得准”，评估基准解决“能比较”。",
    )
    add_table(
        doc,
        ["前沿线", "代表工作", "研究进展", "对你的启发"],
        [
            ("长视频理解", "LongVideoBench、Video-MME、LongVU、StreamingVLM。", "模型开始能处理更长视频，但长上下文成本高、关键信息定位仍难。", "不要直接追大模型训练；应做低成本索引与证据定位。"),
            ("Video-RAG", "Video-RAG、VideoRAG、VideoStir、DrVideo。", "主流思路是先把视频变成文本、事件、帧证据或图结构，再检索相关证据给模型。", "你可复现轻量版：抽帧 + OCR/ASR/检测标签 + 混合检索。"),
            ("多 Agent 视频理解", "HM-RAG、LVAgent、ReAgent-V。", "Agent 被用于任务分解、检索、证据汇总、反思、异常解释。", "你可做工程版多 Agent，不必训练 Agent 策略。"),
            ("在线/流式视频理解", "VideoChat-Online、StreamBridge、LiveStar、StreamingVLM。", "研究开始关注持续视频流、增量记忆、何时主动响应、低延迟。", "这与视频流工作最贴近，但难度高于离线长视频。"),
            ("视频库检索问答", "VLQA、RAVEN、V-Agent。", "从单视频问答扩展到多视频集合、实体发现、schema 生成和跨视频搜索。", "适合做就业作品，但论文范围要控制。"),
            ("垂直场景应用", "异常检测、媒体生产质检、机器人/具身场景、智慧城市事件分析。", "前沿论文和产业都在寻找可落地场景，但数据和评价常是瓶颈。", "你的论文可选择小场景闭环，不必覆盖全行业。"),
        ],
        [1500, 2600, 3000, 2260],
    )
    add_callout(
        doc,
        "非全硕士现实判断",
        "你能做的不是“提出一个新的视频大模型”，而是“把现有开源模型、视频工具和 Agent 编排成一个可评估系统”。只要问题定义清楚、实验对比合理、系统能跑、指标可量化，应用型硕士论文是有空间的。",
        GREEN_FILL,
    )


def add_reproducibility(doc):
    add_heading(doc, "2. 你是否能复现，能复现到什么程度", 1)
    add_table(
        doc,
        ["研究类型", "复现难度", "训练条件", "你能做的版本", "不建议做的版本"],
        [
            ("Video-RAG 轻量管线", "中低", "不需要训练大模型；普通显卡或云 API 即可；需要视频抽帧、OCR/ASR/检测和向量检索。", "复现 Video-RAG/DrVideo 的思想：把视频转为辅助文本和事件索引，再做检索问答。", "不要尝试完整复现论文所有 benchmark 和大模型组合。"),
            ("长视频基准评测", "中", "需要调用若干开源或 API 模型；需要整理问题集和评分脚本。", "选 1-2 个公开数据集 + 自建小数据，评估你的检索 Agent 是否优于 naive RAG。", "不要做大规模 benchmark 排行榜。"),
            ("多 Agent 协作", "中", "通常不需要训练，但需要良好的工具接口、状态管理和日志。", "做检索 Agent、证据 Agent、回答 Agent、评估 Agent 的轻量协作。", "不要做强化学习或复杂策略训练。"),
            ("在线流式视频理解", "中高", "需要视频流接入、增量缓存、低延迟推理；最好有 GPU 或云模型。", "做低频抽帧 + 滑动窗口记忆 + 主动事件摘要。", "不要训练 StreamingVLM 或 LiveStar 这类模型。"),
            ("视频大模型训练/微调", "高", "需要大量视频数据、GPU、训练经验和评测资源。", "只做 LoRA/提示/检索增强层面的轻量实验。", "不建议作为主论文路线。"),
            ("图结构/实体发现", "中高", "不一定训练，但需要 schema、实体抽取、图存储和复杂评估。", "小规模视频库实体/事件图，作为进阶创新点。", "不要做跨大规模视频库的完整知识图谱系统。"),
        ],
        [1500, 1200, 2500, 2600, 1560],
    )
    add_para(
        doc,
        "对你来说，最稳的技术路线是：FFmpeg/GStreamer 或 OpenCV 做视频接入与抽帧，OCR/ASR/检测模型生成辅助文本，SQLite/PostgreSQL/向量库做索引，OpenAI Agents SDK 或 LangGraph 做工具调用和 Agent 流程，最后用 Recall@K、时间定位误差、问答正确率、证据命中率和延迟做评估。",
    )


def add_innovation_strategy(doc):
    add_heading(doc, "3. 非全背景下的创新策略", 1)
    add_para(
        doc,
        "你的创新不应放在“模型参数层”，而应放在“系统方法层”。这不是降低要求，而是换一个适合你的发力点：你有工程背景、视频业务入口和 C++ 系统感，适合把多个已有能力组合成可验证的新系统。",
    )
    add_table(
        doc,
        ["创新层级", "可行性", "论文可写性", "示例"],
        [
            ("事件 schema 创新", "高", "较强。可以明确写方法。", "把视频片段组织为时间段、证据帧、对象、语音、文字、动作、场景、置信度、来源。"),
            ("混合检索策略", "高", "较强。可做消融实验。", "关键词 + 向量 + 时间窗口 + 事件类型 + 证据优先级。"),
            ("Agent 工具链设计", "高", "中强。要避免只写工程实现。", "让 Agent 显式调用 search_events、get_evidence、verify_answer、generate_report。"),
            ("在线记忆机制", "中", "较强。更前沿但实现复杂。", "滑动窗口、事件记忆、长期摘要、主动触发。"),
            ("垂直场景数据集", "中", "较强。适合应用型论文。", "自建小规模长视频问题集，标注事件和证据帧。"),
            ("模型训练创新", "低", "强但不适合你当前条件。", "训练新视频模型、强化学习 Agent、复杂 VLM 微调。"),
        ],
        [1800, 1300, 1800, 4460],
    )
    add_callout(
        doc,
        "推荐创新组合",
        "最适合你的组合是：事件 schema + 混合检索 + Agent 工具链 + 小规模可控评估。这个组合不需要大规模训练，但能形成论文方法、系统实现、实验对比和求职作品。",
        BLUE_FILL,
    )


def add_five_topics(doc):
    add_heading(doc, "4. 五种视频流 + Agent 论文题目方案", 1)
    add_para(
        doc,
        "下面 5 个题目都属于“视频流 + Agent”方向，但重心不同。它们按你的背景、时间、论文可写性和就业相关性进行了设计。建议 6 月先做 2 个最小验证，8 月前收束为 1 个正式开题题目。",
    )
    add_table(
        doc,
        ["编号", "题目方向", "核心问题", "推荐度"],
        [
            ("A", "面向长时视频流的语义事件索引与 Agent 检索问答系统", "如何把长视频转成可检索事件，并让 Agent 返回带证据的答案。", "最高：论文与就业兼顾，训练成本低。"),
            ("B", "面向实时视频流的增量事件记忆与主动摘要 Agent", "视频流持续进入时，Agent 如何边看边记、按需总结、主动提示。", "高：更前沿，但实现复杂度更高。"),
            ("C", "面向视频库的多模态实体/事件图谱构建与 Agent 查询系统", "多视频集合中如何自动抽取实体、事件和关系，支持跨视频查询。", "中高：创新性强，但数据和评估更重。"),
            ("D", "面向视频系统工程的画面-日志联合诊断 Agent", "如何结合视频画面、播放日志、推拉流状态定位故障或异常。", "高：最贴近你的工作和就业，但视觉学术性要补。"),
            ("E", "面向垂直场景的轻量 Video-RAG Agent 事件解释系统", "在园区、会议、机器人或教学等场景中，如何低成本做事件解释。", "中高：应用性强，关键在选好场景。"),
        ],
        [800, 3500, 3300, 1760],
    )

    add_heading(doc, "5. 五个题目的技术要求、训练条件与创新性对比", 1)
    add_table(
        doc,
        ["方向", "技术要求", "训练条件", "创新性", "适合度"],
        [
            ("A 语义事件索引 QA", "视频抽帧、OCR/ASR/检测标签、事件 schema、向量/时间混合检索、Agent 工具调用。", "基本不需要训练；可用开源模型或 API；需要小规模标注评估集。", "中高。创新在事件结构、混合检索、证据可追溯和评估。", "最适合。可 9 月开题，风险可控。"),
            ("B 实时流记忆 Agent", "RTSP/RTMP/WebRTC 或本地视频流、滑动窗口、增量摘要、事件记忆、主动触发策略。", "不训练也可做；若追求效果需使用视频 LLM/API；硬件和延迟要求更高。", "高。在线理解和主动响应是前沿问题。", "适合但需降级实现，先做低频抽帧版。"),
            ("C 视频库实体图谱 Agent", "多视频管理、实体抽取、schema 生成、图数据库/关系索引、跨视频检索。", "不需要训练，但需要较多标注和数据清洗。", "高。与 RAVEN/VideoRAG 接近，可扩展性强。", "中等。适合后续扩展，不建议第一个做。"),
            ("D 画面-日志联合诊断 Agent", "视频播放/推流日志、码率/延迟/丢帧指标、异常帧采样、Agent 诊断流程。", "不需要模型训练；需要构造故障样例和日志数据。", "中。学术前沿性弱一些，但工程价值和就业价值强。", "很适合做备选或论文工程增强模块。"),
            ("E 垂直场景 Video-RAG", "根据场景选择事件类型、辅助文本、检索策略和问答模板。", "不训练或少量微调；核心是数据集和场景定义。", "中高。创新依赖场景选择和评估设计。", "适合。若导师偏应用场景，可作为开题题目。"),
        ],
        [1700, 3100, 1900, 2200, 460],
    )


def add_topic_details(doc):
    add_heading(doc, "6. 五个题目的详细分析", 1)

    topics = [
        (
            "A. 面向长时视频流的语义事件索引与 Agent 检索问答系统",
            "这是当前最推荐题目。它把长视频理解问题拆成视频切片、辅助文本抽取、事件索引、混合检索和 Agent 问答。论文可以对比直接帧输入、纯文本 RAG、仅 OCR/ASR/检测标签、无 Agent 工具链等 baseline。就业上可对应 AI 视频工程、Agent 工程和多模态应用工程。",
            "风险是题目容易写大。建议限定为 3-5 小时公开视频或自建视频，50-100 个标注事件，50-100 个问题，重点做证据可追溯和检索评估。",
        ),
        (
            "B. 面向实时视频流的增量事件记忆与主动摘要 Agent",
            "这是更前沿的方向，贴近在线视频理解、流式 VLM 和实时 Agent。它的核心是视频流持续进入时，系统如何维护短期窗口、长期事件记忆，并在满足条件时主动摘要或告警。",
            "风险是实时性和工程复杂度高。适合做降级版：每 2-5 秒抽帧，按 30-60 秒窗口生成事件摘要，Agent 只在触发规则或用户查询时工作。",
        ),
        (
            "C. 面向视频库的多模态实体/事件图谱构建与 Agent 查询系统",
            "这个题目把单视频扩展到多视频集合，强调实体、事件、关系和 schema。它适合连接 RAVEN、VideoRAG 等工作，也有视频资产管理、媒体内容检索、教育课程视频、会议视频库等职业场景。",
            "风险是数据整理和图谱评估成本较高。非全背景下建议只做小规模视频库，不做大规模实体发现。",
        ),
        (
            "D. 面向视频系统工程的画面-日志联合诊断 Agent",
            "这个题目最贴近你的当前工作。它把视频画面异常、播放状态、推拉流日志、码率、延迟、丢帧、插件日志放到一起，让 Agent 生成诊断建议和证据链。",
            "风险是导师可能认为视觉研究含量不足。解决办法是在方法里加入异常帧采样、画面质量检测、事件定位和多模态证据融合，让它不只是日志分析。",
        ),
        (
            "E. 面向垂直场景的轻量 Video-RAG Agent 事件解释系统",
            "这个题目把方法放进一个场景，例如园区巡检、会议视频、课堂视频、机器人第一视角、设备操作教学等。优点是应用叙事清晰，便于开题和答辩。",
            "风险是场景选择决定成败。不要选择强隐私、强安防或需要大量真实数据的场景；优先选择公开视频或可自采样例。",
        ),
    ]
    for title, value, risk in topics:
        add_heading(doc, title, 2)
        add_para(doc, value)
        add_callout(doc, "主要风险与收束方式", risk, CAUTION_FILL)


def add_recommendation(doc):
    add_heading(doc, "7. 推荐排序与 6-9 月验证计划", 1)
    add_table(
        doc,
        ["排序", "题目", "为什么"],
        [
            ("1", "A 语义事件索引 + Agent QA", "最平衡：可复现、可开题、可就业、训练成本低。"),
            ("2", "D 画面-日志联合诊断 Agent", "最贴近工作：可转化为工程作品，但要增强视觉与多模态部分。"),
            ("3", "B 实时流记忆 Agent", "最前沿：适合做亮点，但应作为 A 的进阶版本或子模块。"),
            ("4", "E 垂直场景 Video-RAG", "取决于数据场景；若导师喜欢应用场景，可提升优先级。"),
            ("5", "C 视频库实体图谱 Agent", "创新性强但工程和数据成本较高，适合作为后续拓展。"),
        ],
        [1000, 3300, 5060],
    )
    add_table(
        doc,
        ["时间", "动作", "目标"],
        [
            ("6 月", "同时做 A 和 D 的最小原型：A 做视频事件索引，D 做日志/画面联合诊断草图。", "判断哪个更能和工作、导师方向、个人兴趣结合。"),
            ("7 月", "读 8-10 篇核心论文，复现 Video-RAG/DrVideo 的轻量思想。", "证明方向不是空想，形成开题相关工作。"),
            ("8 月", "做一个可演示系统：视频输入、事件抽取、索引、Agent 查询、证据返回。", "让导师看到可行性和论文边界。"),
            ("9 月", "正式开题：优先 A，允许加入 B 的在线记忆或 D 的工程诊断作为特色。", "题目既前沿又可完成。"),
        ],
        [1200, 5600, 2560],
    )
    add_callout(
        doc,
        "最终建议",
        "不要把自己锁死在单一题目上。6-8 月的目标是验证“视频流 + Agent”这个方向中哪一个题目最适合你。正式开题时，建议以 A 为主线，选择 B 或 D 的一个模块作为创新增强点。",
        BLUE_FILL,
    )


def add_papers(doc):
    add_heading(doc, "8. 关键论文与资料", 1)
    add_table(
        doc,
        ["资料", "主要内容", "阅读价值"],
        [
            ("DrVideo, CVPR 2025", "把长视频理解转化为文档检索问题，先构造视频文档，再检索关键帧和文本证据。", "证明“视频转文档/事件索引”是顶会认可方向。"),
            ("Video-RAG, arXiv 2024", "训练无关、低成本，使用音频、OCR、目标检测等辅助文本增强长视频理解。", "最适合你复现轻量版。"),
            ("VideoRAG, arXiv 2025", "面向超长视频，结合文本知识 grounding、多模态上下文和跨视频知识图。", "提供进阶上限，不必完整复现。"),
            ("VideoStir, arXiv 2026", "结构化、意图感知的长视频 RAG，指出普通 RAG 会破坏时空结构和隐式意图。", "支持你做时序结构和查询意图。"),
            ("HM-RAG, Pattern Recognition 2026", "层次化多 Agent RAG，用时间记忆、外部知识和多源融合做长视频推理与异常检测。", "多 Agent 论文对标。"),
            ("LongVideoBench, NeurIPS 2024", "长上下文视频语言理解基准，视频可到约一小时，强调引用推理。", "实验指标和问题设计参考。"),
            ("Video-MME, CVPR 2025", "综合评估多模态大模型视频分析能力，覆盖多任务、多时长、多领域。", "开题中证明方向前沿。"),
            ("VideoChat-Online / OVBench, CVPR 2025", "面向在线空间-时间视频理解，评估流式上下文中的视频理解能力。", "B 方向的核心参考。"),
            ("RAVEN, arXiv 2025", "Agentic 框架，用于大规模视频集合的实体发现、schema 生成和结构化表示。", "C 方向的核心参考。"),
            ("OpenAI Agents SDK / LangGraph", "工具调用、工作流、状态、handoff、trace、评估与部署。", "Agent 工程实现参考。"),
        ],
        [2200, 4100, 3060],
    )
    sources = [
        "DrVideo: https://openaccess.thecvf.com/content/CVPR2025/html/Ma_DrVideo_Document_Retrieval_Based_Long_Video_Understanding_CVPR_2025_paper.html",
        "Video-RAG: https://arxiv.org/abs/2411.13093",
        "VideoRAG: https://arxiv.org/abs/2502.01549",
        "VideoStir: https://arxiv.org/abs/2604.05418",
        "HM-RAG: https://www.sciencedirect.com/science/article/pii/S003132032600587X",
        "LongVideoBench: https://proceedings.neurips.cc/paper_files/paper/2024/hash/329ad516cf7a6ac306f29882e9c77558-Abstract-Datasets_and_Benchmarks_Track.html",
        "Video-MME: https://openaccess.thecvf.com/content/CVPR2025/html/Fu_Video-MME_The_First-Ever_Comprehensive_Evaluation_Benchmark_of_Multi-modal_LLMs_in_CVPR_2025_paper.html",
        "VideoChat-Online / OVBench: https://papers.cool/venue/Huang_Online_Video_Understanding_OVBench_and_VideoChat-Online%40CVPR2025%40CVF",
        "RAVEN: https://arxiv.org/abs/2504.06272",
        "OpenAI Agents SDK: https://developers.openai.com/api/docs/guides/agents",
        "LangGraph: https://docs.langchain.com/oss/python/langgraph/workflows-agents",
    ]
    for source in sources:
        add_bullet(doc, source)


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_doc()
    add_cover(doc)
    add_frontier_map(doc)
    add_reproducibility(doc)
    add_innovation_strategy(doc)
    add_five_topics(doc)
    add_topic_details(doc)
    add_recommendation(doc)
    add_papers(doc)
    doc.core_properties.title = "Video Stream + Agent Thesis Topic Space Report"
    doc.core_properties.subject = "Feasibility and topic comparison for video stream plus agent thesis directions"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    build()

