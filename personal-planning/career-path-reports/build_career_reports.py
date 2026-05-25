from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
REPORT_DATE = "2026-05-22"
FONT_BODY = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 45, 60)
MUTED = RGBColor(92, 104, 116)
FILL_GRAY = "F2F4F7"
FILL_BLUE = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size: float | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text: str = "", style: str | None = None, bold: bool = False, color=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold, color=color)
    return p


def add_body(doc, text: str):
    p = add_paragraph(doc, text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    return p


def add_bullet(doc, text: str):
    p = add_paragraph(doc, text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p


def add_callout(doc, label: str, text: str, fill: str = CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(label)
    set_run_font(run, bold=True, color=DARK_BLUE)
    run = p.add_run("：" + text)
    set_run_font(run, color=INK)
    doc.add_paragraph()
    return table


def add_kv_table(doc, rows: list[tuple[str, str]], header: tuple[str, str] | None = None):
    table = doc.add_table(rows=1 if header else 0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    if header:
        for idx, text in enumerate(header):
            cell = table.cell(0, idx)
            set_cell_shading(cell, FILL_GRAY)
            run = cell.paragraphs[0].add_run(text)
            set_run_font(run, bold=True, color=INK)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], FILL_BLUE)
        run = cells[0].paragraphs[0].add_run(key)
        set_run_font(run, bold=True, color=INK)
        run = cells[1].paragraphs[0].add_run(value)
        set_run_font(run, color=INK)
    doc.add_paragraph()
    return table


def add_score_table(doc, scores: dict[str, tuple[int, str]]):
    rows = [
        ("兴趣匹配", "interest_fit"),
        ("学习成本", "learning_cost"),
        ("成都可行性", "chengdu_fit"),
        ("薪资上限", "salary_ceiling"),
        ("稳定性", "stability"),
        ("AI 抗冲击", "ai_resilience"),
        ("转型风险", "transition_risk"),
        ("个人精力匹配", "energy_fit"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, [2100, 1500, 5760])
    headers = ["维度", "评分", "解释"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, FILL_GRAY)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, bold=True, color=INK)
    for label, key in rows:
        score, note = scores[key]
        cells = table.add_row().cells
        values = [label, f"{score}/5", note]
        for idx, value in enumerate(values):
            run = cells[idx].paragraphs[0].add_run(value)
            set_run_font(run, color=INK, bold=(idx == 0))
            if idx == 1:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table


def add_plan_table(doc, rows: list[tuple[str, str, str]]):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, [1600, 3920, 3840])
    headers = ["阶段", "验证动作", "通过信号"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, FILL_GRAY)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, bold=True, color=INK)
    for stage, action, signal in rows:
        cells = table.add_row().cells
        for idx, value in enumerate([stage, action, signal]):
            run = cells[idx].paragraphs[0].add_run(value)
            set_run_font(run, color=INK, bold=(idx == 0))
    doc.add_paragraph()
    return table


def add_curve_table(doc, rows: list[tuple[str, str, str]]):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, [1600, 3940, 3820])
    headers = ["时间", "合理目标", "判断重点"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, FILL_GRAY)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, bold=True, color=INK)
    for t, goal, check in rows:
        cells = table.add_row().cells
        for idx, value in enumerate([t, goal, check]):
            run = cells[idx].paragraphs[0].add_run(value)
            set_run_font(run, color=INK, bold=(idx == 0))
    doc.add_paragraph()
    return table


COMMON_SOURCES = [
    "WEF《The Future of Jobs Report 2025》：技术变化、AI、技能重塑会持续影响 2025-2030 年劳动力市场。https://www.weforum.org/publications/the-future-of-jobs-report-2025/",
    "BLS Occupational Outlook Handbook：软件开发、质量保障与测试岗位 2024-2034 年预计增长 15%，但传统 Computer Programmer 岗位预计下降 6%，重复编码任务受自动化影响更明显。https://www.bls.gov/ooh/computer-and-information-technology/software-developers.htm",
    "GitHub Octoverse 2024：生成式 AI 项目、Python、Jupyter Notebook、AI agent 和小模型生态快速增长，说明 AI 原生开发正在变成软件工程的一部分。https://github.blog/news-insights/octoverse/octoverse-2024/",
    "四川省人民政府网站：成都元宇宙行动方案提出 2025 年相关产业规模目标 1500 亿元；成都具身智能、机器人、AI 实景验证和数字文创场景正在扩展。https://www.sc.gov.cn/10462/10464/10465/10595/2023/1/4/fdcc2f41865a4202a332d2903b500098.shtml",
    "教育部等五部门《关于进一步做好非全日制研究生就业工作的通知》：全日制和非全日制研究生学历学位证书具有同等法律地位和相同效力，但具体报考仍需以岗位公告为准。https://www.moe.gov.cn/srcsite/A22/s7065/202002/t20200214_421068.html",
]


REPORTS = [
    {
        "filename": "xr-3d-vision-path-report.docx",
        "title": "XR / 3D Vision Development Path Report",
        "subtitle": "AR、VR、三维视觉、空间计算与数字孪生方向决策报告",
        "verdict": "值得作为兴趣优先的高潜力探索方向，但不建议马上把它当成唯一主线。它最吸引你的地方是未来感、空间交互和视觉系统；最大问题是学习曲线陡、岗位密度有限、成都机会需要更细致筛选。",
        "route_definition": [
            "这条路包括 AR/VR/MR 交互、空间计算、三维重建、数字孪生、点云/深度相机应用、Unity/Unreal 可视化、城市级 3D 场景和工业/文旅/教育中的沉浸式体验。",
            "它不等于传统监控运维，也不应被压缩成单纯工厂质检。你真正感兴趣的是“人如何在三维空间里理解和操作数字世界”，而不是反复调摄像头或看告警。",
            "更现实的切入点不是纯算法论文，而是 3D 应用工程、视觉数据处理、空间交互原型、数字孪生工具链、C++/Python 与引擎侧能力的结合。",
        ],
        "fit": [
            "这条路与 INTP 式好奇心高度匹配：它有空间、视觉、交互、系统结构，也有足够多的新概念可探索。",
            "它的风险也很清楚：三维方向会遇到线性代数、坐标变换、相机模型、渲染管线、SLAM/重建等内容。如果只靠热情冲刺，容易在数学和工程细节里消耗掉。",
            "你的优势不在于成为最强 3D 算法研究员，而在于把视觉、视频、C++ 工程和产品场景连接起来，做能跑、能展示、能解释的空间视觉系统。",
        ],
        "ai_value": "AI 会加强这条路，而不是简单替代它。生成式 3D、NeRF/Gaussian Splatting、视觉大模型、空间理解、多模态交互会降低原型成本，但真实系统仍需要工程集成、实时性能、传感器数据、交互体验和场景落地。你的价值应放在“把 AI 视觉能力接入三维空间应用”。",
        "chengdu": "成都有元宇宙、数字文创、游戏、科幻文旅、城市数字孪生和智慧城市线索。四川省政府网站披露，成都元宇宙行动方案曾提出到 2025 年相关产业规模达到 1500 亿元，并提到三维实景城市、物理/游戏引擎输出、灵境导航等场景。问题是，高质量 3D 视觉岗位不一定密集，岗位可能分散在数字文创、城市空间、游戏引擎、测绘遥感、智慧城市、车载/机器人等公司里。",
        "income": "成都 20k 目标有机会，但通常要求你不是“入门 3D 学习者”，而是能独立做空间视觉/引擎工具/算法工程化的人。若只会 Unity 展示或简单建模，薪资天花板会偏低；若能结合 C++、OpenCV、点云、实时视频、引擎和 AI 推理，性价比会明显提高。",
        "learning": "以你每周约 22 小时的投入，3 个月可以验证兴趣，6 个月可以做出作品，但 12 个月内想达到可跳槽水平，需要控制学习范围。建议先做“3D 视觉应用工程”而不是大而全学图形学、SLAM、引擎、深度学习。",
        "scores": {
            "interest_fit": (5, "未来感强，与你对 AR/VR、3D、空间技术的兴趣高度一致。"),
            "learning_cost": (2, "数学、图形、相机模型和引擎都要补，不能低估。"),
            "chengdu_fit": (3, "成都有政策和场景，但岗位密度与一线城市相比需要验证。"),
            "salary_ceiling": (4, "做深后上限不错，泛展示型岗位上限一般。"),
            "stability": (3, "行业波动比传统软件大，项目制和政策场景较多。"),
            "ai_resilience": (4, "AI 会降低创作门槛，但真实空间系统仍需要工程落地。"),
            "transition_risk": (3, "与你当前视频/C++有桥，但仍需要补较多新栈。"),
            "energy_fit": (3, "适合探索，但必须避免贪多导致疲惫。"),
        },
        "curve": [
            ("1 年", "完成 OpenCV + 基础相机模型 + Unity/Unreal 之一 + 一个空间视觉 Demo。", "是否能持续投入，而不是只停留在兴趣幻想。"),
            ("3 年", "形成 3D 应用工程/空间视觉工程师定位，可进入数字孪生、XR 应用或空间数据公司。", "是否有可迁移作品和成都岗位入口。"),
            ("5 年", "向空间计算系统、XR 技术负责人、数字孪生解决方案架构方向发展。", "是否仍有行业红利，以及个人是否愿意长期追技术栈变化。"),
        ],
        "risks": [
            "被“酷技术”吸引，但长期卡在数学和底层图形细节。",
            "成都岗位可能更多偏展示、项目交付、政府/文旅场景，技术深度不一定稳定。",
            "如果进入强项目制公司，可能出现阶段性加班；需要避开长期混乱型团队。",
            "过早追 SLAM/NeRF/Gaussian Splatting 论文，可能导致执行成本失控。",
        ],
        "validation": [
            ("30 天", "做一个 RGB-D/双目/单目视频的空间理解小实验，输出英文技术笔记。", "你愿意在坐标、相机、可视化这些细节上继续花时间。"),
            ("90 天", "完成一个 Unity/Unreal 或 Python 可视化 Demo：视频/图像输入，输出 3D 场景、目标位置或空间标注。", "能做出可展示作品，而不是只看课程。"),
            ("180 天", "筛选成都/远程/一线城市 50 个岗位，反推技能缺口；尝试投递或约 3 次行业访谈。", "发现真实岗位与你兴趣重叠，且 20k 路径不是幻想。"),
        ],
        "signals": [
            "增强信号：你能连续做出 3D 视觉 Demo，并且对空间坐标和引擎调试不排斥。",
            "削弱信号：你只喜欢概念，不喜欢基础数学和工程细节。",
            "策略建议：作为主兴趣方向保留，但短期最好与 C++ 视频或 Agent 工具结合，形成更稳的复合能力。",
        ],
    },
    {
        "filename": "robotics-embodied-perception-path-report.docx",
        "title": "Robotics And Embodied Perception Path Report",
        "subtitle": "机器人感知、具身智能、SLAM 相邻方向与智能硬件决策报告",
        "verdict": "值得认真探索，但要避开“听起来是机器人，实际是工厂现场调设备”的路线。更适合你的切入点是机器人感知应用、传感器数据处理、视频/视觉系统集成、具身智能应用工程，而不是纯控制、纯硬件或长期驻场交付。",
        "route_definition": [
            "这条路包括机器人视觉感知、多传感器融合、SLAM 相邻应用、机器人数据采集与评估、具身智能应用、智能硬件视觉模块、服务机器人/文旅机器人/康养机器人等真实场景。",
            "它不等于传统工厂自动化流水线，也不等于长期驻厂调 PLC、调相机、陪产线。那些方向可以赚钱或稳定，但与你的未来感兴趣点不完全一致。",
            "更合适的定位是“懂 C++/视频/视觉的机器人感知工程师或应用工程师”，把摄像头、深度数据、模型推理、机器人场景和产品迭代连接起来。",
        ],
        "fit": [
            "你本科阶段接触过机器人和视觉比赛，这说明机器人不是凭空冒出来的兴趣，而是曾经给你打开过视野的领域。",
            "这条路符合你喜欢未来方向、系统结构和真实产品的特点。机器人是软硬件结合系统，能带来比普通业务开发更强的意义感。",
            "但机器人方向也会考验耐心：调试周期长，硬件约束多，很多问题不是代码写完就结束。如果团队不成熟，混乱感会很强。",
        ],
        "ai_value": "AI 正在把机器人从规则控制推向大模型、视觉语言模型、多模态感知和具身智能。成都近年的公开报道也显示具身智能、机器人实景验证、文旅/康养/警务等场景在推进。但真正落地仍需要工程师理解传感器、实时系统、失败模式和安全边界。",
        "chengdu": "成都机器人和具身智能近两年信号较强。四川省政府网站报道，成都机器人产品已在交通、治安、教育、文旅等城市真实场景验证，也提到 AI 核心产业规模和企业数量突破“双千”。这说明成都不是完全没有机会，但岗位可能更多集中于应用、场景、集成、测试验证和产品化，不一定都是前沿算法岗。",
        "income": "若进入成熟机器人/智能硬件企业，成都 20k 有现实可能，但通常要求能独立处理视觉感知、系统调试、部署和工程问题。若岗位变成长期驻场实施或低端自动化，收入和成长都会打折。",
        "learning": "每周 22 小时投入足以做机器人方向的初步验证。建议先补 ROS2 基础、相机/深度传感器、OpenCV、ONNX 推理和一个仿真环境，不要一上来啃完整 SLAM 或控制理论。",
        "scores": {
            "interest_fit": (4, "未来感强，且与本科机器人/视觉经历有情感连接。"),
            "learning_cost": (2, "涉及传感器、坐标、ROS、部署、硬件调试，成本较高。"),
            "chengdu_fit": (4, "成都具身智能和机器人场景有明确政策和产业信号。"),
            "salary_ceiling": (4, "复合型机器人感知工程能力有较好上限。"),
            "stability": (3, "行业成长快，但公司差异和项目波动较大。"),
            "ai_resilience": (5, "AI 是机器人升级核心变量，工程落地难被纯代码生成替代。"),
            "transition_risk": (3, "与嵌入式/C++/视觉有桥，但软硬件复杂度会增加。"),
            "energy_fit": (3, "兴趣能支撑，但调试挫败感需要提前适应。"),
        },
        "curve": [
            ("1 年", "完成 ROS2 + 摄像头/视频输入 + 目标检测/深度估计 + 简单导航或交互 Demo。", "判断自己是否能接受机器人系统调试。"),
            ("3 年", "进入机器人、智能硬件、具身智能应用企业，负责视觉感知或场景集成。", "避免滑向低端驻场交付。"),
            ("5 年", "成为机器人感知系统负责人、具身智能应用架构师或智能硬件产品技术骨干。", "能否跨软件、传感器、产品场景做系统判断。"),
        ],
        "risks": [
            "实际岗位可能没有想象中未来，反而是现场交付、设备调试和客户催促。",
            "SLAM 和控制如果学太深，会进入你不热爱的复杂数学区。",
            "硬件调试不可控，容易消耗耐心；需要选择工程文化成熟的团队。",
            "成都机会虽有增长，但好岗位数量仍需逐个验证。",
        ],
        "validation": [
            ("30 天", "跑通 ROS2 基础、相机输入和一个现成视觉模型；写一篇学习记录。", "你不排斥机器人开发环境和调试方式。"),
            ("90 天", "做一个小型具身感知 Demo：摄像头识别目标，输出坐标/语义，并驱动仿真或简单动作。", "你能把视频、模型和机器人动作串起来。"),
            ("180 天", "调研成都机器人/智能硬件/具身智能企业与岗位，筛选非驻场、非强销售、非纯交付机会。", "存在与你兴趣和薪资目标同时匹配的岗位池。"),
        ],
        "signals": [
            "增强信号：你喜欢调试传感器和系统行为，不只是喜欢机器人概念。",
            "削弱信号：你无法忍受硬件不稳定和长链路排错。",
            "策略建议：作为未来感技术主线候选，但要用 Demo 和岗位调研过滤掉低端自动化路线。",
        ],
    },
    {
        "filename": "cpp-video-media-ai-path-report.docx",
        "title": "C++ Video Media AI Path Report",
        "subtitle": "C++ 视频系统、实时媒体、插件架构与 AI 推理集成决策报告",
        "verdict": "这是连续性最强、短期性价比最高的技术主线。它未必是最让你兴奋的方向，但它能把你已有的嵌入式、C++、视频插件、导师视觉资源和 AI 工程化连接起来，是适合做保底主线的路线。",
        "route_definition": [
            "这条路包括 C++ 视频处理、FFmpeg/GStreamer/WebRTC、编解码、实时流媒体、插件架构、跨平台客户端、视频 AI 推理接入、视频质量/性能分析工具。",
            "它不等于传统监控运维，也不等于长期看平台告警。你应避免把自己放在低成长的监控项目维护里，而是向视频基础设施、实时媒体系统和 AI 视频能力集成靠拢。",
            "这条路可以向 XR、机器人、Agent 工具继续分叉：XR 需要视频和渲染，机器人需要视觉流，Agent 需要自动化处理视频数据和工程流程。",
        ],
        "fit": [
            "这条路最符合你的现有资产：嵌入式训练让你理解底层约束，C++ 工作让你有工程入口，视频插件让你接触真实业务，导师视觉方向提供研究连接。",
            "它对数学要求相对可控，更多考验工程严谨性、系统理解、性能意识、架构和文档能力。你对代码质量、流程、架构和文档的敏感度在这条路上是优势。",
            "缺点是它的未来感不如 XR/机器人明显，如果长期只做公司内部插件，会有被业务细节困住的风险。",
        ],
        "ai_value": "AI 会改变视频系统：模型推理、视频摘要、质量检测、智能剪辑、异常识别、检索和多模态理解都需要接入视频管线。AI 可能替代部分普通编码，但不容易替代实时性能、工程集成、跨平台兼容和视频系统调试能力。",
        "chengdu": "成都有数字文创、游戏、元宇宙、智慧城市、轨道交通、机器人和 AI 场景，这些都可能需要视频处理或实时媒体能力。但纯高端视频底层岗位不一定多，可能分散在音视频云、安防、会议、文旅互动、数字孪生、轨交智能化、机器人企业中。",
        "income": "成都 20k 目标在这条路上相对现实，前提是你从业务 C++ 开发升级为“视频系统工程师”：熟悉 FFmpeg/GStreamer/WebRTC 至少一个主栈，能做性能分析、架构拆分、AI 推理接入和稳定性治理。",
        "learning": "每周 22 小时投入非常适合这条路。它不需要一次性换赛道，可以边工作边补 FFmpeg/GStreamer/OpenCV/ONNX Runtime，并把工作中的混乱项目整理成架构笔记和工具。",
        "scores": {
            "interest_fit": (3, "不是最兴奋，但与视频、视觉、工程系统有持续连接。"),
            "learning_cost": (4, "已有 C++ 和视频工作基础，补课成本相对可控。"),
            "chengdu_fit": (3, "岗位存在但分散，需要找视频系统而非传统运维。"),
            "salary_ceiling": (4, "成熟视频系统工程师有较强议价空间。"),
            "stability": (4, "视频基础设施和工程能力迁移性较好。"),
            "ai_resilience": (4, "AI 需要视频系统承载，工程化能力抗冲击较强。"),
            "transition_risk": (4, "与现有经历连续，切换风险最低。"),
            "energy_fit": (4, "适合在工作和研究之间稳步积累。"),
        },
        "curve": [
            ("1 年", "建立 FFmpeg/GStreamer/WebRTC/OpenCV/ONNX 中 2-3 项核心能力，完成一个 AI 视频处理 Demo。", "能否从业务插件走向通用视频系统能力。"),
            ("3 年", "成为视频系统/实时媒体/AI 视频工程主力，具备架构、性能、稳定性经验。", "是否拿到成都或远程 20k 级机会。"),
            ("5 年", "向视频系统架构师、AI 视频平台负责人、XR/机器人视觉系统工程专家延展。", "是否持续接触更有未来感的上层场景。"),
        ],
        "risks": [
            "长期停留在公司内部插件维护，能力不可迁移。",
            "被吸入传统监控、运维和低质量项目，兴趣被消耗。",
            "只写 C++ 业务代码，不补视频链路和 AI 推理，AI 时代竞争力会变弱。",
            "当前团队工程水平不足，可能让你养成妥协习惯。",
        ],
        "validation": [
            ("30 天", "整理当前视频插件架构，补一篇 FFmpeg/GStreamer/OpenCV 之一的英文技术笔记。", "你能把工作经验抽象成可迁移知识。"),
            ("90 天", "做一个本地视频 AI Demo：读取视频流，做检测/分割/摘要/质量分析，再输出结果。", "你能把视频系统和 AI 能力串起来。"),
            ("180 天", "完成一个可展示项目和简历重写，调研成都音视频/数字文创/机器人/智慧城市岗位。", "这条路能支撑 20k 成都目标并保留未来分叉。"),
        ],
        "signals": [
            "增强信号：你能在现有工作中主动抽象架构和工具，产生作品。",
            "削弱信号：你发现自己对视频底层完全无兴趣，只想做更上层产品或 AI 应用。",
            "策略建议：作为短中期主线最稳，配合 XR/机器人/Agent 作为兴趣分叉。",
        ],
    },
    {
        "filename": "agent-development-engineer-path-report.docx",
        "title": "Agent Development Engineer Path Report",
        "subtitle": "AI Agent、LLM 应用工程、RAG、工具调用与工作流自动化决策报告",
        "verdict": "这是最值得新增为第六条的高上行路线。它比纯算法研究更适合你，也比传统程序员更贴近 AI 浪潮。但它竞争会迅速变卷，必须尽早形成“工程系统 + 场景理解 + 工具链落地”的差异化，而不是只会调用 API。",
        "route_definition": [
            "这条路包括 LLM 应用开发、Agent 工作流、工具调用、RAG、知识库、代码/文档/数据自动化、多 Agent 协作、评估体系、权限与安全、AI 原生产品工程。",
            "它不等于简单套壳聊天机器人，也不等于只会 LangChain 教程。真正有价值的是让 Agent 连接真实工具、业务流程、代码仓库、文档、数据库、视频数据和用户目标。",
            "你的特别切入点可以是“面向视频/视觉/C++工程的 Agent 开发”：例如自动读论文、分析视频日志、生成插件文档、辅助代码评审、调度视频处理流程。",
        ],
        "fit": [
            "这条路非常适合你的好奇心和系统思维。Agent 开发不是单点算法，而是理解任务、工具、上下文、边界、评估和产品体验。",
            "它允许你绕开一部分复杂数学，把重点放在工程、流程、结构化思考和快速实验上。你的 INTP 式探索欲在这里可能变成优势。",
            "但它也有一个风险：变化太快，容易被新框架、新模型、新概念牵着走。如果没有真实场景和作品，很容易变成浅层追热点。",
        ],
        "ai_value": "GitHub Octoverse 2024 显示生成式 AI 项目和 AI agent 相关探索在快速增长，Python 与 Jupyter 的增长也说明 AI 工程正在吸引大量开发者。AI Agent 不是“远离程序员”，而是把程序员从写局部代码转向设计智能工作流、工具接口、评估和自动化系统。",
        "chengdu": "成都的 Agent 岗位密度可能不如北京、上海、深圳、杭州，但成都有 AI、数字政府、文创、教育、机器人、智慧城市和企业数字化场景。Agent 能作为跨行业能力迁移，未来也可能通过远程、外包转产品、内部工具或创业小产品实现成都生活。",
        "income": "成都 20k 目标在 Agent 路线中有机会，但前提是你不能只做提示词或 API 调用。你需要证明自己能做后端工程、工具集成、RAG 质量治理、Agent 评估、部署、成本控制和真实业务闭环。若能结合 C++ 视频/视觉，差异化会更强。",
        "learning": "每周 22 小时足够高强度推进。与 XR/机器人相比，Agent 路线见效更快：30 天可做原型，90 天可做可用工具，180 天可以形成作品集。关键是每个项目都要解决真实问题。",
        "scores": {
            "interest_fit": (4, "未来感强、反馈快、适合探索和系统组合。"),
            "learning_cost": (4, "入门成本低于 XR/机器人，但深入评估、架构和产品化仍有门槛。"),
            "chengdu_fit": (3, "本地岗位需验证，但远程和跨行业迁移性较好。"),
            "salary_ceiling": (4, "AI 原生工程有上行空间，差异化越强上限越高。"),
            "stability": (3, "技术变化快，岗位定义会频繁调整。"),
            "ai_resilience": (5, "直接站在 AI 工具化和自动化建设一侧。"),
            "transition_risk": (3, "需要从 C++ 转向 Python/TypeScript/后端/产品工程，但作品验证快。"),
            "energy_fit": (5, "适合短周期实验和持续新鲜感。"),
        },
        "curve": [
            ("1 年", "掌握 Python/TypeScript 后端、RAG、工具调用、Agent 编排、评估和部署，完成 3 个真实工具。", "是否能从玩模型转为解决真实问题。"),
            ("3 年", "成为 AI 应用/Agent 工程师，能独立负责企业内部工具、知识库、自动化流程或 AI 产品模块。", "是否形成行业场景差异化。"),
            ("5 年", "向 AI 产品技术负责人、AI 平台工程师、Agent 系统架构或独立工具产品方向发展。", "是否具备产品判断和工程稳定性能力。"),
        ],
        "risks": [
            "赛道拥挤，很多人都会写 Demo；没有业务场景就容易同质化。",
            "框架变化快，追工具名词会消耗精力。",
            "如果只做前端聊天界面，薪资和壁垒都有限。",
            "需要补 Web 后端、部署、数据、权限、安全和评估，这些不比写算法轻松。",
        ],
        "validation": [
            ("30 天", "做一个个人规划/论文阅读 Agent：能读取 Markdown、总结论文、生成行动清单。", "你喜欢设计工具流程，而不是只和模型聊天。"),
            ("90 天", "做一个 C++/视频项目 Agent：读取代码、日志或文档，输出架构摘要、问题清单和测试建议。", "你能把旧经验转化为 AI 工具差异化。"),
            ("180 天", "完成 2-3 个可展示 Agent 项目，调研成都/远程 AI 应用岗位，准备转型简历。", "Agent 路线能支撑 20k 成都目标或远程机会。"),
        ],
        "signals": [
            "增强信号：你能快速做出工具，并持续改进评估和可用性。",
            "削弱信号：你只喜欢追模型新闻，不愿处理工程细节、接口和部署。",
            "策略建议：这是最适合作为新增重点探索的路线，可与 C++ 视频、论文分析、个人知识库结合起来。",
        ],
    },
    {
        "filename": "ai-product-transition-path-report.docx",
        "title": "AI Product Transition Path Report",
        "subtitle": "AI 产品经理、技术产品、解决方案与产品运营转型决策报告",
        "verdict": "可以作为大转型候选，但必须谨慎区分“技术产品”与“强销售/纯协调/背锅型产品”。你愿意大转型，这是机会；但你的优势来自工程理解和系统判断，不应放弃技术核去做低门槛运营。",
        "route_definition": [
            "这条路包括 AI 产品经理、技术产品经理、解决方案产品、内部工具产品、AI 平台产品、视频/视觉/机器人相关产品规划。",
            "它不适合走强销售、纯商务、纯运营、长期陪客户、无技术含量的需求传话筒路线。",
            "最适合你的版本是“懂工程的 AI 产品/技术产品”：能理解模型边界、系统架构、成本、数据、交付风险和用户场景。",
        ],
        "fit": [
            "你愿意大转型，并且对工程质量、流程和架构有判断力，这些都是技术产品的好基础。",
            "INTP 倾向可能让你擅长抽象系统、拆问题、找结构，但产品岗会要求更强的沟通、推进、妥协和对人性的耐心。",
            "你讨厌强销售和长期结构性加班，因此不能去客户压力极重、交付混乱、老板拍脑袋的产品岗。",
        ],
        "ai_value": "AI 产品岗的价值在于判断哪些需求可以用 AI 做、怎样评估效果、怎样控制成本、如何把模型能力嵌入流程。AI 会降低部分原型门槛，但不会替代产品判断、场景取舍、跨团队沟通和落地责任。",
        "chengdu": "成都有数字政府、文创、教育、游戏、智慧城市、AI、机器人和企业数字化场景，技术产品岗位可能存在于国企数字化平台、AI 应用公司、ToB 软件公司、文旅科技和智能硬件企业。但高质量产品岗通常要求行业经验和沟通成果，不只是技术背景。",
        "income": "成都 20k 对产品路线可达，但通常要求你能独立负责模块、拿到明确结果，或者有 AI/视频/工业/政企解决方案经验。普通产品运营或初级 PM 可能低于这个目标。",
        "learning": "每周 22 小时可用于补产品方法、行业调研、竞品分析、PRD、原型和访谈。相比技术路线，它不难入门，但难在真实项目经验和人际推进能力。",
        "scores": {
            "interest_fit": (3, "如果能做未来技术产品会有兴趣，纯协调会厌倦。"),
            "learning_cost": (4, "知识入门较快，但能力验证依赖真实协作场景。"),
            "chengdu_fit": (4, "成都政企、文创、AI 应用和数字化公司有机会。"),
            "salary_ceiling": (4, "技术产品做深有上限，普通运营上限较低。"),
            "stability": (3, "产品岗受公司质量、老板风格和业务成败影响大。"),
            "ai_resilience": (4, "AI 产品化能力会更重要，但低端产品文档岗会被压缩。"),
            "transition_risk": (2, "从工程转产品意味着评价体系和日常工作方式大变。"),
            "energy_fit": (3, "适合思考和结构化，不一定适合高频扯皮。"),
        },
        "curve": [
            ("1 年", "从技术产品/AI 应用产品助理型角色切入，完成 2-3 份高质量 PRD/竞品/方案作品。", "你是否喜欢沟通推进而不只是分析。"),
            ("3 年", "成为 AI/视频/机器人/企业数字化方向产品经理，能负责独立模块。", "是否能拿到结果和业务信任。"),
            ("5 年", "向技术产品负责人、解决方案负责人、AI 产品 Lead 或业务负责人发展。", "是否愿意承担更强的人和业务压力。"),
        ],
        "risks": [
            "产品岗可能变成协调、催进度、写材料和背锅，而不是思考未来技术。",
            "强销售和客户陪伴型岗位与你偏好冲突。",
            "脱离技术太快会失去差异化。",
            "如果表达和推进能力没有刻意训练，转型初期会痛苦。",
        ],
        "validation": [
            ("30 天", "选一个 Agent/视频/机器人产品，写一份竞品分析和需求拆解。", "你是否享受从用户、场景、商业角度思考。"),
            ("90 天", "把一个个人 Agent 或视频工具做成完整 PRD、原型和技术可行性说明。", "你能否把技术想法转成产品语言。"),
            ("180 天", "尝试内部推动一个小工具或找技术产品岗位访谈，验证真实工作日常。", "你能接受沟通、推进和不确定性。"),
        ],
        "signals": [
            "增强信号：你能把复杂技术讲清楚，并愿意和非技术人反复沟通。",
            "削弱信号：你讨厌会议、催办、需求变化和人际协调。",
            "策略建议：不要直接跳纯产品，先走 AI 技术产品/解决方案产品，保留技术核。",
        ],
    },
    {
        "filename": "chengdu-public-sector-stable-path-report.docx",
        "title": "Chengdu Public Sector Stable Path Report",
        "subtitle": "成都公务员、事业编、国企信息化与稳定路线决策报告",
        "verdict": "可以作为严肃备选，但不建议盲目押注一年。你的行为模式显示：开始阶段会有强动机，但长期重复复习、背诵和高压考试会使动力衰减。因此这条路需要阶段性闸门，而不是简单说“我愿意准备一年”。",
        "route_definition": [
            "这条路包括四川/成都公务员、事业单位、国企数字化、智慧城市、信息中心、政务信息化、网络安全、数据治理、公共服务平台等稳定岗位。",
            "它不等于完全躺平，也不等于放弃技术。更现实的版本是用非全硕士、计科背景和工作经验争取对口信息化岗位。",
            "它也不应与技术路线完全二选一：可以先用 3-6 个月做岗位调研和考试适应性测试，再决定是否投入完整一年。",
        ],
        "fit": [
            "你对稳定生活和成都有真实偏好，且能接受公务员约 1w 月薪，这说明体制路线不是纯逃避，而是有生活锚点。",
            "但你的历史模式需要认真看：小升初、中考、高考、考研都不是优势项；考研初期热情高，后期在重复复习和背诵中动力下降。这意味着公务员/事业编备考存在较高行为风险。",
            "如果准备方式仍是“每天刷题背资料，熬一年等考试”，成功概率不乐观；如果改成外部监督、月度模考、岗位筛选、阶段淘汰，成功率会提高。",
        ],
        "ai_value": "AI 对体制岗位的冲击相对慢，但数字政府、数据治理、智能办公、政务 AI 应用会增加对复合型信息化人才的需求。你的技术背景可能在信息中心、数据平台、网络安全、智慧城市等岗位更有用。",
        "chengdu": "这是成都落地性最强的路线之一。非全硕士在教育部文件中具有同等法律地位和相同效力，公务员、事业单位、国企公开招聘原则上应提供平等机会；但具体岗位仍要看公告是否要求学历、学位、专业、基层经历、应届身份或其他条件。",
        "income": "公务员或事业编约 1w 月薪你可以接受，但要综合公积金、福利、稳定性、晋升速度、隐性压力和工作内容。国企信息化或平台公司可能比编制收入高，但稳定性和加班情况差异较大。",
        "learning": "你的每周 22 小时足以备考，但关键不是时间，而是重复性承受能力。建议把一年拆成 3 个闸门：第 1 个月看是否能稳定刷题，第 3 个月看模考排名，第 6 个月决定是否全力押注。",
        "scores": {
            "interest_fit": (2, "生活匹配高，但工作内容未必激发兴趣。"),
            "learning_cost": (2, "备考内容重复、背诵和应试强，踩中你的弱点。"),
            "chengdu_fit": (5, "成都定居和稳定性匹配度最高。"),
            "salary_ceiling": (2, "编制收入上限有限，但生活确定性强。"),
            "stability": (5, "稳定性是核心优势。"),
            "ai_resilience": (4, "短中期抗冲击强，信息化岗位仍有需求。"),
            "transition_risk": (3, "一旦上岸稳定，但备考失败会消耗一年机会成本。"),
            "energy_fit": (2, "长期重复复习与你行为模式冲突。"),
        },
        "curve": [
            ("1 年", "完成岗位筛选、备考、模考和考试；同时保留技术作品最低更新频率。", "能否坚持重复训练且模考进入可竞争区间。"),
            ("3 年", "若上岸，适应体制规则并争取信息化/数据/项目管理相关职责。若未上岸，回到技术或产品路线。", "是否获得稳定生活而不是长期压抑。"),
            ("5 年", "在成都形成稳定生活，向信息化骨干、项目负责人或数字政府相关岗位积累。", "是否接受收入天花板和组织规则。"),
        ],
        "risks": [
            "一年备考可能重复考研失败模式：初期热情高，后期厌倦。",
            "岗位公告可能限制专业、学历形式、应届身份或基层经验，不能只看政策原则。",
            "上岸后工作可能并不轻松，也可能有材料、会议、协调和隐性加班。",
            "如果完全停止技术积累，未上岸后的回撤成本会很高。",
        ],
        "validation": [
            ("30 天", "收集近三年四川/成都公务员、事业单位、国企信息化岗位公告，筛出 20 个理论可报岗位。", "岗位真实存在，非全硕士和专业条件不构成明显阻碍。"),
            ("90 天", "按公考节奏完成行测/申论基础轮和 3 次模拟考，记录分数曲线和厌倦程度。", "分数上升且没有明显心理排斥。"),
            ("180 天", "若模考进入可竞争区间，继续全力；若长期停滞，转为备选并恢复技术/Agent 主线。", "用数据决定，不用焦虑决定。"),
        ],
        "signals": [
            "增强信号：你能连续 3 个月稳定刷题，模考分数持续提升，且能接受材料写作。",
            "削弱信号：你在第 2 个月就开始明显逃避重复题和背诵。",
            "策略建议：作为稳定路线保留，但必须设置月度闸门，避免把一年押成新的考研循环。",
        ],
    },
]


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Title", 20, BLUE),
        ("Subtitle", 12, MUTED),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name != "Subtitle"

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167

    return doc


def add_cover(doc: Document, report: dict) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(report["title"])
    set_run_font(run, size=20, bold=True, color=BLUE)
    title.paragraph_format.space_after = Pt(3)

    subtitle = doc.add_paragraph(style="Subtitle")
    run = subtitle.add_run(report["subtitle"])
    set_run_font(run, size=12, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(12)

    add_kv_table(
        doc,
        [
            ("报告用途", "用于比较该发展曲线是否值得投入 3-6 个月验证，而不是直接决定最终职业目标。"),
            ("个人基准", "成都技术岗 20k RMB/月为满意线；公务员/事业编约 10k RMB/月可接受。"),
            ("精力假设", "工作日每天约 2 小时，周末每天约 6 小时，每周约 22 小时可用于学习和项目。"),
            ("偏好约束", "偏好机器人、三维、AI 等未来方向；不喜欢强销售；不接受长期结构性加班。"),
            ("日期", REPORT_DATE),
        ],
    )
    add_callout(doc, "一句话判断", report["verdict"])


def add_report(doc: Document, report: dict) -> None:
    add_cover(doc, report)

    add_heading(doc, "1. 路线定义", 1)
    for item in report["route_definition"]:
        add_bullet(doc, item)

    add_heading(doc, "2. 与你的适配度", 1)
    for item in report["fit"]:
        add_body(doc, item)

    add_heading(doc, "3. AI 时代价值", 1)
    add_body(doc, report["ai_value"])

    add_heading(doc, "4. 成都可行性与收入生活匹配", 1)
    add_body(doc, report["chengdu"])
    add_body(doc, report["income"])

    add_heading(doc, "5. 学习成本与可执行性", 1)
    add_body(doc, report["learning"])

    add_heading(doc, "6. 评分矩阵", 1)
    add_body(doc, "评分为 1-5 分。高分不代表一定选择，低分也不代表彻底排除；它用于暴露这条路的真实代价。")
    add_score_table(doc, report["scores"])

    add_heading(doc, "7. 1 年 / 3 年 / 5 年发展曲线", 1)
    add_curve_table(doc, report["curve"])

    add_heading(doc, "8. 风险和陷阱", 1)
    for item in report["risks"]:
        add_bullet(doc, item)

    add_heading(doc, "9. 30 / 90 / 180 天验证计划", 1)
    add_plan_table(doc, report["validation"])

    add_heading(doc, "10. 决策信号", 1)
    for item in report["signals"]:
        add_bullet(doc, item)

    add_heading(doc, "资料依据", 1)
    for source in COMMON_SOURCES:
        add_bullet(doc, source)

    add_callout(
        doc,
        "使用建议",
        "先把这份报告当作路线地图，不要把它当成结论。下一步应选择 1-2 条最有兴趣的路径做 30 天验证，用真实产出和情绪反馈修正判断。",
        FILL_BLUE,
    )


def add_header_footer(doc: Document, title: str) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(title)
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"Personal Planning Decision Report | {REPORT_DATE}")
    set_run_font(run, size=9, color=MUTED)


def save_report(report: dict) -> Path:
    doc = setup_document()
    add_header_footer(doc, report["title"])
    add_report(doc, report)
    path = OUTPUT_DIR / report["filename"]
    doc.save(path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    generated = [save_report(report) for report in REPORTS]
    for path in generated:
        print(path)


if __name__ == "__main__":
    main()
