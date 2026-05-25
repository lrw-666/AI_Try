from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = OUTPUT_DIR / "cpp-video-agent-development-roadmap-2026-2028.docx"
REPORT_DATE = "2026-05-22"

FONT_BODY = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 45, 60)
MUTED = RGBColor(92, 104, 116)
GRAY_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF4D6"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text="", style=None, size=11, bold=False, color=INK, after=6, line=1.25):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run_font(r, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    p = add_para(doc, text, style="List Bullet", after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    return p


def add_number(doc, text):
    p = add_para(doc, text, style="List Number", after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    return p


def add_callout(doc, label, text, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ": ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths, header_fill=GRAY_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            if idx == 0 and len(str(value)) <= 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, color=INK, bold=(idx == 0))
    doc.add_paragraph()
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Title", 22, BLUE),
        ("Subtitle", 12, MUTED),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("C++ Video + Agent Roadmap")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Personal Career Development Plan | {REPORT_DATE}")
    set_run_font(r, size=9, color=MUTED)
    return doc


def add_cover(doc):
    p = doc.add_paragraph(style="Title")
    r = p.add_run("C++ Video + Agent Development Roadmap 2026-2028")
    set_run_font(r, size=22, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(style="Subtitle")
    r = p.add_run("从视频边缘业务走向视频流数据化、Agent 工作流与硕士论文落地")
    set_run_font(r, size=12, color=MUTED)
    p.paragraph_format.space_after = Pt(12)

    add_table(
        doc,
        ["项目", "内容"],
        [
            ("规划周期", "2026 年 6 月至 2028 年 6 月，覆盖开题、研究、项目、就业与毕业。"),
            ("核心定位", "以 C++ 视频系统为底座，向视频流数据化、多模态检索分析、Agent 工具链扩展。"),
            ("个人约束", "工作日每天约 2 小时，周末每天约 6 小时；偏好未来感技术，不喜欢长期结构性加班和强销售。"),
            ("毕业目标", "完成可答辩的应用型硕士论文，同时形成可投递成都 20k 技术岗位的作品组合。"),
            ("推荐论文线", "面向长时视频流的语义事件索引与 Agent 检索问答系统。"),
        ],
        [1900, 7460],
        BLUE_FILL,
    )
    add_callout(
        doc,
        "核心判断",
        "你不需要硬闯 XR 或机器人最深的技术栈，也不需要退回传统监控运维。更适合的路是把视频播放经验扩展为视频流处理、视频数据抽取、事件索引、AI 分析和 Agent 工作流。这样你既能接触 XR、机器人、智慧城市、数字孪生等业务，又能保留自身 C++/视频工程的可迁移壁垒。",
        CAUTION_FILL,
    )


def add_executive_summary(doc):
    add_heading(doc, "1. 总体定位", 1)
    add_para(
        doc,
        "这份规划把你未来两年的主线定义为“C++ 视频系统 + AI Agent 工程化”。它不是简单从视频播放转向监控，也不是直接跳入 XR/机器人/SLAM 的深水区，而是把视频作为数据入口，把流媒体管线作为工程底座，把多模态模型和 Agent 作为能力放大器。",
    )
    add_para(
        doc,
        "这条路的关键好处是连续性强：你当前做视频插件，过去有 C++ 和嵌入式训练，导师在视觉领域，未来又希望接触机器人、3D、Agent 等更有未来感的方向。视频流正好是这些方向的共同输入层。XR 需要视频和渲染，机器人需要视觉流和传感器数据，Agent 需要能调用工具处理视频、日志、文档和事件。",
    )
    add_callout(
        doc,
        "两年目标",
        "到 2028 年 6 月毕业时，你应当不只是“会写 C++ 的视频插件工程师”，而是能展示一套“视频流接入、事件抽取、语义索引、Agent 查询、工程化部署”的完整系统，并能把它讲成论文、项目和求职故事。",
    )
    add_table(
        doc,
        ["层级", "能力含义", "对应产出"],
        [
            ("底座层", "C++、视频插件、FFmpeg/GStreamer、RTSP/RTMP/WebRTC、ONVIF、推拉流、性能与稳定性。", "视频流实验仓库、协议笔记、推拉流 Demo、性能测试记录。"),
            ("数据层", "抽帧、关键帧、OCR/ASR、目标检测、事件切片、元数据结构化、向量索引。", "视频事件数据库、索引 schema、可重复评估脚本。"),
            ("智能层", "多模态模型、Video-RAG、Agent 工具调用、工作流编排、检索问答、异常解释。", "视频问答 Agent、事件检索 Agent、论文实验原型。"),
            ("职业层", "把视频系统能力接入 XR、机器人、智慧城市、数字孪生、AI 平台或企业内部工具。", "简历项目、技术博客、论文、成都岗位匹配清单。"),
        ],
        [1500, 4860, 3000],
    )


def add_strategy(doc):
    add_heading(doc, "2. 为什么这条路适合你", 1)
    add_bullet(doc, "它保留了你已有资产：C++、嵌入式训练、视频插件经验、视觉导师资源，以及对工程质量的敏感度。")
    add_bullet(doc, "它绕开了你不喜欢的部分：纯工厂流水线检测、传统监控运维、强销售、长期无意义加班、纯数学公式推导。")
    add_bullet(doc, "它给你未来感入口：视频流可以进入机器人视觉、XR 空间理解、数字孪生、智慧城市、工业安全和 Agent 工作流。")
    add_bullet(doc, "它适合你的行为模式：用可见项目驱动学习，而不是长期重复刷题；每个月都能看到 Demo、笔记或论文进展。")
    add_bullet(doc, "它适合非全硕士节奏：可以把工作场景、论文方向和个人作品统一起来，减少三线分裂。")

    add_heading(doc, "3. 目标岗位画像", 1)
    add_table(
        doc,
        ["目标岗位", "你要证明的能力", "成都 20k 可能性"],
        [
            ("AI 视频系统工程师", "能处理视频流、编解码、推拉流、AI 推理接入、事件抽取和系统稳定性。", "较高，但要求不是只会播放插件，而是能做完整视频处理链路。"),
            ("多模态 Agent 工程师", "能让 Agent 调用工具处理视频、文档、日志、数据库和模型结果，并能做评估与权限控制。", "中高，本地岗位需验证；远程或一线跳板可增强机会。"),
            ("视频平台/流媒体工程师", "熟悉 FFmpeg、GStreamer、WebRTC、RTSP/RTMP、媒体服务器和跨平台性能。", "中高，岗位分散在音视频、文创、智慧城市、会议、机器人和数字孪生企业。"),
            ("视觉应用工程师", "能把检测、跟踪、OCR、事件识别等模型接入业务，不追求纯算法原创。", "中等，需避开低端工厂质检和纯驻场交付。"),
            ("技术产品/解决方案", "能把视频+AI+Agent 系统讲成产品方案、交付路径和客户价值。", "中等到较高，但必须避开强销售和纯协调角色。"),
        ],
        [2100, 4900, 2360],
    )


def add_learning_roadmap(doc):
    add_heading(doc, "4. 学习路线：四层能力栈", 1)
    add_para(doc, "学习路线不建议按“课程列表”推进，而应按工程层级推进。每一层都必须有可运行产出，否则容易陷入 INTP 式概念兴奋但缺少落地。")
    add_table(
        doc,
        ["层级", "优先学习内容", "不优先学习内容", "最低产出"],
        [
            ("视频基础", "FFmpeg 命令与 libavformat/libavcodec 概念；RTSP/RTMP/HLS/WebRTC 基础；编码、封装、帧率、时间戳、同步。", "深入写编码器、完整实现媒体服务器内核。", "能从 RTSP 拉流、转码、切片、推送，并记录延迟、丢帧和 CPU 占用。"),
            ("流媒体工程", "GStreamer pipeline 思维；ONVIF Profile T；媒体服务器使用；推拉流链路；日志与监控指标。", "一开始就啃完整 WebRTC C++ Native 栈。", "搭建本地多路视频接入、转推和截图/抽帧服务。"),
            ("视频数据化", "抽帧策略、关键帧、目标检测、OCR/ASR、事件切片、元数据 schema、向量数据库。", "追求最前沿视频大模型训练。", "把一段长视频转为可检索事件表和向量索引。"),
            ("Agent 工程", "工具调用、RAG、工作流编排、状态管理、human-in-the-loop、观测与评估、安全边界。", "只做聊天机器人 UI 或追框架名词。", "Agent 能调用视频检索工具回答“何时发生了什么、证据帧在哪里”。"),
            ("论文工程", "问题定义、相关工作、实验设计、指标、消融实验、系统架构图、论文写作。", "追求顶会级算法创新。", "形成可答辩的应用型研究：方法清楚、系统可跑、实验可复现。"),
        ],
        [1500, 3200, 2500, 2160],
    )

    add_heading(doc, "5. 推荐技术栈", 1)
    add_table(
        doc,
        ["类别", "推荐选择", "选择理由"],
        [
            ("语言", "C++ 负责视频核心链路；Python 负责 AI/Agent/实验；TypeScript 可选用于前端或 Agent 服务。", "保留 C++ 护城河，同时用 Python 加速 AI 与论文实验。"),
            ("视频", "FFmpeg 必学；GStreamer 二选一深入；WebRTC 先理解概念和应用层；ONVIF 了解设备接入。", "覆盖从文件、网络流、设备、转码到实时通信的主干能力。"),
            ("AI/CV", "OpenCV、ONNX Runtime、PyTorch 基础、目标检测/跟踪/OCR/ASR、CLIP/多模态 embedding。", "足够支撑视频数据化与论文实验，不需要一开始做大模型训练。"),
            ("Agent", "OpenAI Agents SDK 或 LangGraph 作为主框架之一；FastAPI 做服务；SQLite/PostgreSQL + 向量库做索引。", "既能快速原型，也能体现工程化思维。"),
            ("评估", "检索 Recall@K、事件定位误差、问答准确率、延迟、资源消耗、人工标注一致性。", "论文需要可量化指标，项目需要可解释质量。"),
        ],
        [1500, 5000, 2860],
    )


def add_timeline(doc):
    add_heading(doc, "6. 2026.06-2028.06 总时间线", 1)
    add_callout(
        doc,
        "节奏原则",
        "2026 年 6-9 月必须服务开题；2026 年 10 月到 2027 年 6 月必须跑出论文原型；2027 年 7-12 月必须补实验和就业作品；2028 年 1-6 月必须收束论文、答辩和求职。",
    )
    add_table(
        doc,
        ["阶段", "时间", "主目标", "必须产出"],
        [
            ("开题准备", "2026.06-2026.09", "把视频流+Agent 方向收束成可开题问题。", "3 个候选题、15 篇论文笔记、1 个最小 Demo、开题 PPT 和任务书。"),
            ("原型搭建", "2026.10-2027.03", "完成视频接入、事件抽取、索引和 Agent 查询原型。", "系统 v0.1、实验数据集、baseline、初步指标。"),
            ("实验强化", "2027.04-2027.09", "把系统从 Demo 推到论文可评估状态。", "系统 v0.2、消融实验、论文中期材料、1-2 个可展示项目。"),
            ("职业转化", "2027.10-2028.01", "把论文系统转化为简历项目和成都岗位竞争力。", "简历、作品集、岗位地图、投递策略、论文初稿 60%。"),
            ("毕业冲刺", "2028.02-2028.06", "完成论文、答辩、求职或回成都落地。", "论文终稿、答辩材料、项目文档、求职结果或备选路线。"),
        ],
        [1600, 1700, 3600, 2460],
    )

    monthly_rows = [
        ("2026.06", "方向收束", "复盘当前视频插件业务，画出播放链路、数据链路、日志链路；学习 FFmpeg 基础命令和 RTSP/RTMP/HLS 概念。", "列出 5 个论文题目草案，和导师约一次方向沟通。"),
        ("2026.07", "视频流基础", "跑通本地 RTSP 拉流、转码、截图、抽帧；整理推流/拉流/转码实验笔记。", "阅读 5 篇视频理解/Video-RAG/多模态检索论文。"),
        ("2026.08", "Agent 原型", "用 Python/FastAPI 做一个视频元数据服务；Agent 能调用工具查询视频片段和帧。", "形成 3 个开题候选题，每个包含问题、方法、数据、指标、风险。"),
        ("2026.09", "开题", "确定论文题目、技术路线、实验数据和最低可行系统。", "完成开题报告、开题 PPT、论文大纲 v0。"),
        ("2026.10", "视频数据化 v0", "实现抽帧、关键帧保存、OCR/ASR/目标检测中的至少 1-2 项。", "建立事件表 schema 和第一版样例数据集。"),
        ("2026.11", "索引与检索", "加入向量索引和时间索引，支持按文本查询视频事件。", "完成 baseline：直接帧检索、文本检索、简单 RAG 对比。"),
        ("2026.12", "Agent 工作流", "Agent 支持检索、证据帧返回、回答生成、失败提示和日志记录。", "系统 v0.1 演示，写阶段总结。"),
        ("2027.01", "实验设计", "确定评价指标：Recall@K、事件定位误差、回答准确率、延迟、资源消耗。", "做 1 次小规模人工标注，验证数据可用性。"),
        ("2027.02", "寒假强化", "集中补 OpenCV、ONNX Runtime、视频抽样策略和 Agent 评估。", "完成论文相关工作初稿。"),
        ("2027.03", "Baseline 完成", "跑通至少 2 个 baseline 和你的方法 v0。", "输出实验表格 v0，和导师确认是否可继续。"),
        ("2027.04", "方法改进", "加入多源辅助文本：OCR、ASR、目标/场景标签、运动事件。", "完成系统 v0.2 架构图和方法章节草稿。"),
        ("2027.05", "中期准备", "补充数据、修正指标、完成可视化界面或报告输出。", "准备中期检查材料。"),
        ("2027.06", "中期检查", "完成一次完整实验闭环，从视频输入到 Agent 输出和指标统计。", "中期报告、系统演示、论文目录定稿。"),
        ("2027.07", "职业映射", "把论文系统拆成简历项目：视频流、索引、Agent、评估、部署。", "整理成都/一线/远程岗位地图 80 个。"),
        ("2027.08", "项目强化", "补一个工程化功能：多路视频、任务队列、日志追踪、权限或成本控制。", "作品集 README、系统部署文档。"),
        ("2027.09", "论文实验扩展", "做消融实验：无 OCR/无 ASR/无 Agent/不同抽帧策略/不同检索策略。", "形成实验章节 v0。"),
        ("2027.10", "求职准备", "根据目标岗位重写简历：AI 视频系统、Agent 工程、视频平台三版。", "开始小范围投递或约行业访谈。"),
        ("2027.11", "论文初稿 40%", "完成绪论、相关工作、系统设计、部分实验。", "导师反馈一轮，确定补实验清单。"),
        ("2027.12", "就业验证", "根据面试反馈修正项目表达；补 C++ 视频系统八股和 Agent 工程案例。", "论文初稿 60%，作品集可公开版本。"),
        ("2028.01", "收束路线", "决定主投成都、先一线跳板、还是继续当前岗位过渡。", "岗位清单、投递记录、论文补实验完成。"),
        ("2028.02", "论文初稿", "春节前后完成论文初稿 80%，集中处理图表、引用、实验复现。", "论文初稿提交导师。"),
        ("2028.03", "答辩预演", "修改论文，准备答辩 PPT；项目演示稳定化。", "模拟答辩 1-2 次。"),
        ("2028.04", "求职冲刺", "集中投递成都与远程岗位；准备谈薪和选择标准。", "拿到面试或 offer 进展；论文终稿接近完成。"),
        ("2028.05", "终稿与答辩", "完成查重、格式、答辩材料和项目备份。", "论文终稿、答辩 PPT、演示视频。"),
        ("2028.06", "毕业转场", "完成毕业答辩，选择就业路线。", "毕业后 3 个月行动计划。"),
    ]
    add_table(
        doc,
        ["月份", "主题", "学习/项目动作", "论文/职业动作"],
        monthly_rows,
        [1200, 1500, 4300, 2360],
    )


def add_thesis_directions(doc):
    add_heading(doc, "7. 论文候选方向", 1)
    add_para(doc, "考虑你 9 月开题、导师视觉方向、自己不想做纯公式算法、当前工作是视频插件，论文应优先选择“应用型研究 + 可跑系统 + 有实验指标”的形态。")
    add_table(
        doc,
        ["方向", "题目草案", "适配度", "主要风险"],
        [
            ("A 推荐", "面向长时视频流的语义事件索引与 Agent 检索问答系统研究", "最高。连接视频流、Agent、Video-RAG、工程系统和导师视觉方向。", "需要控制范围，避免做成泛泛聊天机器人。"),
            ("B 研究更强", "基于多模态辅助文本的长视频检索增强理解方法研究", "较高。更像论文，关注 OCR/ASR/检测标签如何增强长视频理解。", "实验要求更强，需要数据和指标稳定。"),
            ("C 工程更强", "面向视频插件系统的智能日志分析与故障定位 Agent 设计与实现", "中高。贴近工作，可落地。", "视觉含量较弱，需和导师确认是否接受。"),
            ("D 场景更强", "面向园区/城市视频的事件检索与异常解释系统研究", "中等。容易找到应用叙事。", "容易滑向传统监控，需要避免低兴趣场景。"),
            ("E 未来拓展", "面向机器人视觉流的语义记忆与任务查询 Agent 研究", "中等偏高。未来感强。", "机器人数据和平台获取成本较高。"),
        ],
        [1100, 3700, 2600, 1960],
    )

    add_heading(doc, "8. 推荐开题方向详案", 1)
    add_callout(
        doc,
        "推荐题目",
        "面向长时视频流的语义事件索引与 Agent 检索问答系统研究。这个题目足够贴近你的工作，又能接上视觉、多模态、Agent 和未来职业方向；它不要求你发明新模型，但要求你把视频变成可检索、可解释、可问答的结构化知识。",
        BLUE_FILL,
    )
    add_table(
        doc,
        ["模块", "内容"],
        [
            ("研究问题", "长时视频难以直接交给大模型理解，成本高、上下文长、事件稀疏、定位困难。如何低成本抽取视频事件，并让 Agent 按问题检索证据、生成解释？"),
            ("技术路线", "视频流接入 -> 抽帧/切片 -> OCR/ASR/目标/场景标签 -> 时间索引 + 向量索引 -> Agent 工具调用 -> 检索问答/事件摘要 -> 指标评估。"),
            ("创新点", "不是训练大模型，而是把多源辅助文本、时间结构和 Agent 工具调用结合，用工程化方式提升长视频查询效率和可解释性。"),
            ("实验数据", "公开长视频数据集 + 自建小规模业务样例。若真实业务数据敏感，使用开源视频或自行采集无隐私样例。"),
            ("评价指标", "事件检索 Recall@K、时间定位误差、回答准确率、证据帧命中率、平均响应时间、索引构建成本。"),
            ("论文边界", "不追求端到端训练，不做大型视频模型微调，不做完整安防平台，只做“视频语义索引 + Agent 分析”的可复现系统。"),
        ],
        [1800, 7560],
    )

    add_heading(doc, "9. 9 月开题倒排计划", 1)
    add_table(
        doc,
        ["时间", "必须完成", "判断标准"],
        [
            ("6 月第 1-2 周", "把当前工作视频链路画成图：输入、播放、插件、数据、日志、可扩展点。", "能向导师说明你不是空想，而是有工程入口。"),
            ("6 月第 3-4 周", "跑通 FFmpeg 拉流、抽帧、截图、转码；整理实验笔记。", "能证明视频流处理最小闭环可行。"),
            ("7 月", "读 8-10 篇论文，重点 Video-RAG、视频检索、视频异常、Agent 工作流；每篇写英文笔记。", "能形成相关工作地图，而不是只知道几个热词。"),
            ("8 月上旬", "做最小 Demo：上传/接入视频 -> 抽帧 -> 生成标签 -> 建索引 -> Agent 查询。", "能现场演示 3 个问题和证据返回。"),
            ("8 月下旬", "确定题目、研究问题、技术路线、实验指标和风险控制。", "导师认可开题方向，范围不失控。"),
            ("9 月", "完成开题报告和 PPT，准备答辩问答。", "能清楚回答：研究问题是什么、已有工作如何做、你做什么、怎么验证。"),
        ],
        [1700, 5200, 2460],
    )


def add_project_plan(doc):
    add_heading(doc, "10. 项目作品路线", 1)
    add_para(doc, "你的项目不应做成零散 Demo，而应逐步累积成同一个系统的多个版本。最终它既能服务论文，也能服务求职。")
    add_table(
        doc,
        ["版本", "时间", "功能", "求职表达"],
        [
            ("v0.1", "2026.08-2026.12", "单视频输入、抽帧、基础标签、向量索引、Agent 问答。", "我能把视频从播放对象变成可检索数据。"),
            ("v0.2", "2027.01-2027.06", "多源辅助文本、时间索引、事件切片、baseline 和指标。", "我能做视频 AI 系统评估，不只是调模型。"),
            ("v0.3", "2027.07-2027.12", "多路视频、任务队列、日志追踪、权限/成本控制、可视化报告。", "我能做工程化视频 Agent 系统。"),
            ("v1.0", "2028.01-2028.06", "论文稳定版、演示视频、部署文档、简历项目版。", "我能完整讲清问题、架构、指标、取舍和落地价值。"),
        ],
        [1200, 1800, 4400, 1960],
    )

    add_heading(doc, "11. 每周执行系统", 1)
    add_table(
        doc,
        ["时间块", "建议安排", "目的"],
        [
            ("周一 2h", "复盘上周输出，确定本周唯一主任务；写 20 分钟周计划。", "避免方向漂移。"),
            ("周二 2h", "视频工程学习或代码实验。", "补底层能力。"),
            ("周三 2h", "论文阅读或相关工作笔记。", "服务 9 月开题和论文长期积累。"),
            ("周四 2h", "Agent/RAG/后端服务实验。", "保持 AI 原生能力增长。"),
            ("周五 2h", "整理技术笔记、提交代码、记录问题。", "形成可见产出。"),
            ("周六 6h", "长项目块：完成一个可运行功能或实验。", "项目推进主力。"),
            ("周日 6h", "上午补学习，下午写文档/论文/复盘，晚上休息。", "把学习沉淀为论文和作品。"),
        ],
        [1400, 5100, 2860],
    )
    add_callout(
        doc,
        "行为模式提醒",
        "你不适合靠长期重复和意志力硬撑。每周必须有一个看得见的小产出：一张架构图、一段 Demo、一次实验记录、一篇论文笔记或一个可复用脚本。没有产出的学习很容易变成焦虑型收藏。",
        CAUTION_FILL,
    )


def add_job_and_risk(doc):
    add_heading(doc, "12. 就业转化路线", 1)
    add_table(
        doc,
        ["时间", "动作", "目标"],
        [
            ("2026.06-2026.12", "把当前工作从“边缘视频业务”抽象为视频链路、插件架构和数据处理经验。", "简历里不写琐碎业务，写系统能力。"),
            ("2027.01-2027.06", "用论文原型形成第一个作品项目，开始观察成都岗位。", "判断成都 20k 技术岗需要哪些缺口。"),
            ("2027.07-2027.12", "准备三版简历：AI 视频系统、Agent 工程、视频平台工程。", "测试市场反馈，不把自己锁死。"),
            ("2028.01-2028.06", "集中投递成都、远程、一线跳板机会。", "优先成都 20k；若成都机会质量不足，接受短期一线跳板。"),
        ],
        [1700, 5160, 2500],
    )

    add_heading(doc, "13. 风险控制", 1)
    add_table(
        doc,
        ["风险", "表现", "应对策略"],
        [
            ("范围失控", "同时学 FFmpeg、WebRTC、SLAM、3D、Agent、大模型训练。", "只保留视频流 + 索引 + Agent 三件事；XR/机器人只作为业务场景，不作为底层主线。"),
            ("论文不够学术", "系统能跑，但像工程项目，缺少研究问题和指标。", "用长视频理解、Video-RAG、事件检索、问答准确率和消融实验支撑论文性。"),
            ("工作场景受限", "公司只让做边缘业务，接触不到主链路。", "自己用开源工具补链路，形成独立 Demo，不等公司喂机会。"),
            ("Agent 浅层化", "只做聊天界面和提示词。", "必须让 Agent 调用真实工具：视频检索、帧定位、日志查询、报告生成。"),
            ("成都岗位不足", "目标岗位少或薪资不达预期。", "保留一线跳板和远程机会；用论文项目提高跨城市竞争力。"),
            ("持续性下降", "学习两个月后疲惫或逃避。", "每月做一次闸门评估：产出、情绪、导师反馈、市场信号。"),
        ],
        [1600, 3600, 4160],
    )


def add_checkpoints_and_sources(doc):
    add_heading(doc, "14. 月度检查表", 1)
    add_table(
        doc,
        ["检查项", "通过标准"],
        [
            ("学习", "本月至少完成 2 篇论文笔记或 2 篇技术笔记。"),
            ("代码", "本月至少有 1 个可运行功能或实验脚本。"),
            ("论文", "本月至少推进一个论文材料：相关工作、方法图、实验表、导师沟通记录。"),
            ("职业", "每季度更新一次岗位地图、简历或作品集。"),
            ("身体与情绪", "如果连续两周无法推进，不加码自责，改为缩小任务范围。"),
        ],
        [2400, 6960],
    )

    add_heading(doc, "15. 参考资料与依据", 1)
    sources = [
        "OpenAI Agents SDK 文档：Agent 是能够规划、调用工具、跨专家协作并保持状态以完成多步工作的应用；SDK 路线适合开发者自己控制工具、状态、审批和运行时行为。https://developers.openai.com/api/docs/guides/agents",
        "OpenAI《A practical guide to building agents》：Agent 工作流可使用 handoff 等方式在多个专门 Agent 间转移任务，适合复杂流程。https://cdn.openai.com/business-guides-and-resources/a-practical-guide-to-building-agents.pdf",
        "LangGraph 文档：Agent/工作流系统强调工具调用、结构化输出、短期记忆、持久化、流式处理、调试和部署。https://docs.langchain.com/oss/python/langgraph/workflows-agents",
        "FFmpeg 官方文档：FFmpeg 可从文件、网络流、采集设备等多种输入读取，并输出到多种格式，是视频流处理基础工具。https://www.ffmpeg.org/ffmpeg.html",
        "ONVIF Profile T：面向 IP 视频系统，支持 H.264/H.265、成像设置、运动/遮挡告警、元数据流、PTZ 等能力。https://www.onvif.org/profiles/profile-t/",
        "Video-RAG 论文：长视频理解受上下文和计算资源限制，Video-RAG 通过视觉对齐辅助文本和检索增强降低成本，适合作为论文相关工作。https://arxiv.org/abs/2411.13093",
        "四川省人民政府网站：成都 2025 年人工智能产业链目标包括 AI 核心产业规模、算力、大模型、数据集、典型应用场景、机器人产业园等，说明成都存在 AI/机器人/场景应用机会。https://www.sc.gov.cn/10462/10464/10465/10595/2025/3/28/5a9be7787403485083e17b95edae09b0.shtml",
        "四川省人民政府网站：成都元宇宙行动方案提出数字孪生、感知交互、人工智能等核心技术方向，并曾提出 2025 年相关产业规模目标。https://www.sc.gov.cn/10462/10464/10465/10595/2023/1/4/fdcc2f41865a4202a332d2903b500098.shtml",
    ]
    for source in sources:
        add_bullet(doc, source)


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_doc()
    add_cover(doc)
    add_executive_summary(doc)
    add_strategy(doc)
    add_learning_roadmap(doc)
    add_timeline(doc)
    add_thesis_directions(doc)
    add_project_plan(doc)
    add_job_and_risk(doc)
    add_checkpoints_and_sources(doc)
    doc.core_properties.title = "C++ Video + Agent Development Roadmap 2026-2028"
    doc.core_properties.subject = "Personal career, thesis, and project roadmap"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    build()

